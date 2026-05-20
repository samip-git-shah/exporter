import os
import sys
import customtkinter as ctk
import logging
from tkcalendar import DateEntry
import threading
import queue
import re
from playwright.sync_api import sync_playwright
from datetime import datetime
from tkinter import filedialog


def resource_path(rel_path: str) -> str:
    """Resolve a bundled resource path for both source and PyInstaller frozen runs."""
    base = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel_path)


def _log_dir() -> str:
    """Writable directory for logs — alongside source in dev, alongside the exe when frozen."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


logging.basicConfig(filename=os.path.join(_log_dir(), 'app.log'), filemode='w',
                    format='%(name)s - %(levelname)s - %(message)s', level=logging.INFO)

print("Attempting to launch the application with ctk widgets...")

class TestCaseExporter(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.title("Test Case Exporter - By Samip Shah")
        self.state('zoomed')

        # Main container
        main_container = ctk.CTkFrame(self, corner_radius=0)
        main_container.pack(fill="both", expand=True)

        # Left panel
        left_panel = ctk.CTkScrollableFrame(main_container, width=400, corner_radius=10, fg_color="#2c3e50")
        left_panel.pack(side="left", fill="y", padx=10, pady=10)

        # Right panel with preview area
        right_panel = ctk.CTkFrame(main_container, corner_radius=10, fg_color="#34495e")
        right_panel.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        # Preview area header
        preview_header = ctk.CTkLabel(right_panel, text="Follow these instructions to export the Test Case Run from TP",
                                     font=ctk.CTkFont(size=22, weight="bold"))
        preview_header.pack(pady=20)

        # Scrollable preview text area
        self.preview_text = ctk.CTkTextbox(right_panel, wrap="word", font=ctk.CTkFont(size=12),
                                           fg_color="white", text_color="black")
        self.preview_text.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        # Initial instructions text
        instructions_text = """a. Enter the Test Case Run ID of the Test Case Run to be exported

b. Enter your 'Sycamore Email' as the TP User Name

c. Enter your 'Non SSO TP Password' as the TP Password

Note: Before clicking on the 'Get Data from TP' ensure that all the below fields are well populated under the Test Plan Run page:
- Test Machine OS
- Browser Name
- Browser Version
- Host OS
- URL

d. Click on the 'Get Data from TP' button.

   Note: Lookout for the status displayed on top in Green color.

e. Once all the Parameters are loaded click on the:
   • 'Preview' button to Preview the formated Test Case Run in Right Pannel
   • 'Export to Word' to export the formated Test Case Run in Word file
   • 'Export to PDF' to export the formated Test Case Run in PDF file"""

        self.preview_text.insert("1.0", instructions_text)

        # --- Left Panel Content ---

        # Header
        header_label = ctk.CTkLabel(left_panel, text="Test Case Run Exporter", font=ctk.CTkFont(size=24, weight="bold"))
        header_label.pack(pady=(20, 20), padx=20, anchor="w")

        # Status Label
        self.status_label = ctk.CTkLabel(left_panel, text="Ready", font=ctk.CTkFont(size=12, weight="bold"),
                                          text_color="#2ecc71", wraplength=360)
        self.status_label.pack(pady=(0, 10), padx=20, anchor="w")

        # Test Case Run ID
        id_label = ctk.CTkLabel(left_panel, text="Test Case Run ID*", font=ctk.CTkFont(size=14))
        id_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.id_entry = ctk.CTkEntry(left_panel, placeholder_text="Enter ID", corner_radius=10)
        self.id_entry.pack(pady=(0, 10), padx=20, fill="x")
        self.id_entry.bind("<KeyRelease>", lambda e: self.validate_required_fields())

        # Credentials
        user_name_label = ctk.CTkLabel(left_panel, text="TP User Name*", font=ctk.CTkFont(size=14))
        user_name_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.user_name_entry = ctk.CTkEntry(left_panel, placeholder_text="Enter User Name", corner_radius=10)
        self.user_name_entry.pack(pady=(0, 10), padx=20, fill="x")
        self.user_name_entry.bind("<KeyRelease>", lambda e: self.validate_required_fields())

        password_label = ctk.CTkLabel(left_panel, text="TP Password*", font=ctk.CTkFont(size=14))
        password_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.password_entry = ctk.CTkEntry(left_panel, placeholder_text="Enter Password", corner_radius=10, show="*")
        self.password_entry.pack(pady=(0, 20), padx=20, fill="x")
        self.password_entry.bind("<KeyRelease>", lambda e: self.validate_required_fields())

        # Buttons
        self.load_page_button = ctk.CTkButton(left_panel, text="Get Data from TP", corner_radius=10,
                                               fg_color="#3498db", hover_color="#2980b9", command=self.load_page)
        self.load_page_button.pack(pady=10, padx=20, fill="x")
        self.load_page_button.configure(state="disabled")  # Initially disabled

        # IDs (after Load Parameters button)
        test_case_id_label = ctk.CTkLabel(left_panel, text="Test Case ID", font=ctk.CTkFont(size=14))
        test_case_id_label.pack(pady=(20, 5), padx=20, anchor="w")
        self.test_case_id_entry = ctk.CTkEntry(left_panel, placeholder_text="Test Case ID (will be loaded)", corner_radius=10)
        self.test_case_id_entry.pack(pady=(0, 10), padx=20, fill="x")

        test_plan_run_id_label = ctk.CTkLabel(left_panel, text="Test Plan Run ID", font=ctk.CTkFont(size=14))
        test_plan_run_id_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.test_plan_run_id_entry = ctk.CTkEntry(left_panel, placeholder_text="Test Plan Run ID (will be loaded)", corner_radius=10)
        self.test_plan_run_id_entry.pack(pady=(0, 10), padx=20, fill="x")

        # --- Export Parameters Section ---
        session_params_label = ctk.CTkLabel(left_panel, text="Export Parameters", font=ctk.CTkFont(size=16, weight="bold"))
        session_params_label.pack(pady=(30, 10), padx=20, anchor="w")

        test_case_run_name_label = ctk.CTkLabel(left_panel, text="Test Script Name", font=ctk.CTkFont(size=14))
        test_case_run_name_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.test_case_run_name_entry = ctk.CTkEntry(left_panel, placeholder_text="Test Script Name (will be loaded)", corner_radius=10)
        self.test_case_run_name_entry.pack(pady=(0, 10), padx=20, fill="x")

        release_label = ctk.CTkLabel(left_panel, text="Release", font=ctk.CTkFont(size=14))
        release_label.pack(pady=(10, 5), padx=20, anchor="w")
        self.release_entry = ctk.CTkEntry(left_panel, placeholder_text="Release (will be loaded)", corner_radius=10)
        self.release_entry.pack(pady=(0, 10), padx=20, fill="x")

        self.form_entries = {}
        self.date_entries = {}

        # Storage for test case steps data
        self.test_case_steps = []
        self.prerequisites_text = ""
        self.screenshots_text = ""
        self.screenshots_images = []  # Store image data separately
        self.screenshots_content_items = []  # Store ordered content (text and images)

        # Define fields and create the UI elements in a structured way
        # Mapping UI label to the label text on the webpage
        self.field_mapping = {
            "Author Name": "Creator",
            "Last Updated By": "Modified By",
            "Executor Name": "Tester",
            "Services Approver": "Services Reviewer",
            "Product Management Approver": "ProdMgmt Reviewer",
            "Development Approver": "Head of Eng. Reviewer",
            "Execution Status": "Run Result",
            "Test Operating System": "Test Machine OS",
            "Browser Name": "Browser Name",
            "URL": "URL"
        }
        
        self.date_field_mapping = {
            "Author Date": "Creation date",
            "Last Updated Date": "Last Modified",
            "Execution Start Date": "Start Run Date",
            "Execution End Date": "End Run Date"
        }

        # Create text entries
        for label_text in self.field_mapping.keys():
            label = ctk.CTkLabel(left_panel, text=label_text, font=ctk.CTkFont(size=14))
            label.pack(pady=(10, 5), padx=20, anchor="w")
            entry = ctk.CTkEntry(left_panel, placeholder_text=f"...", corner_radius=10)
            entry.pack(pady=(0, 10), padx=20, fill="x")
            self.form_entries[label_text] = entry
            
        # Create date entries with DD-MON-YYYY format
        for label_text in self.date_field_mapping.keys():
            label = ctk.CTkLabel(left_panel, text=label_text, font=ctk.CTkFont(size=14))
            label.pack(pady=(10, 5), padx=20, anchor="w")
            # Use locale 'en_US' with date_pattern 'short' to get a reasonable format
            # We'll format the display ourselves in the code
            date_entry = DateEntry(left_panel, width=15, background='darkblue', foreground='white', borderwidth=2, locale='en_US')
            date_entry.pack(fill="x", padx=20)
            self.date_entries[label_text] = date_entry

        # Action Buttons at the end
        action_buttons_label = ctk.CTkLabel(left_panel, text="Actions", font=ctk.CTkFont(size=16, weight="bold"))
        action_buttons_label.pack(pady=(30, 10), padx=20, anchor="w")

        preview_button = ctk.CTkButton(left_panel, text="Preview", corner_radius=10, fg_color="#9b59b6", hover_color="#8e44ad", command=self.preview_data)
        preview_button.pack(pady=10, padx=20, fill="x")

        self.export_word_button = ctk.CTkButton(left_panel, text="Export to Word", corner_radius=10, fg_color="#e67e22", hover_color="#d35400", command=self.export_to_word)
        self.export_word_button.pack(pady=10, padx=20, fill="x")
        self.export_word_button.configure(state="disabled")  # Initially disabled

        self.export_pdf_button = ctk.CTkButton(left_panel, text="Export to PDF", corner_radius=10, fg_color="#e74c3c", hover_color="#c0392b", command=self.export_to_pdf)
        self.export_pdf_button.pack(pady=10, padx=20, fill="x")
        self.export_pdf_button.configure(state="disabled")  # Initially disabled

        self.cmd_queue = queue.Queue()
        self.playwright_thread = threading.Thread(target=self.playwright_worker, daemon=True)
        self.playwright_thread.start()

    def playwright_worker(self):
        playwright = None
        browser = None
        context = None
        page = None

        def get_field_value(p, label):
            logging.info(f"Attempting to extract value for label: '{label}'")
            try:
                # Wait for content to load
                p.wait_for_timeout(2000)

                # Use JavaScript evaluation for more reliable extraction
                js_code = f"""
                () => {{
                    // Find all elements containing the label text
                    const allElements = Array.from(document.querySelectorAll('*'));

                    // Look for span or div containing exactly the label
                    for (let elem of allElements) {{
                        if (elem.textContent && elem.textContent.trim() === '{label}') {{
                            // Found the label element, now find the value
                            // Try next sibling
                            if (elem.nextElementSibling) {{
                                let value = elem.nextElementSibling.textContent.trim();
                                if (value && value !== '{label}') {{
                                    return value;
                                }}
                            }}

                            // Try parent's next sibling
                            if (elem.parentElement && elem.parentElement.nextElementSibling) {{
                                let value = elem.parentElement.nextElementSibling.textContent.trim();
                                if (value && value !== '{label}') {{
                                    return value;
                                }}
                            }}
                        }}
                    }}

                    // Alternative: Look for label as part of element text
                    for (let elem of allElements) {{
                        const text = elem.textContent || '';
                        if (text.includes('{label}')) {{
                            // Check if this element has the value nearby
                            const children = Array.from(elem.children || []);
                            for (let i = 0; i < children.length; i++) {{
                                if (children[i].textContent.includes('{label}')) {{
                                    // Next sibling might have the value
                                    if (children[i + 1]) {{
                                        let value = children[i + 1].textContent.trim();
                                        if (value && value !== '{label}' && value.length < 200) {{
                                            return value;
                                        }}
                                    }}
                                }}
                            }}
                        }}
                    }}

                    return null;
                }}
                """

                value_text = p.evaluate(js_code)

                if value_text:
                    logging.info(f"Successfully extracted value for '{label}': '{value_text}'")
                    return value_text
                else:
                    logging.warning(f"Could not find value for '{label}' using JavaScript evaluation")
                    return "Not Found"

            except Exception as e:
                logging.error(f"Could not extract value for label '{label}'. Error: {e}")
                return "Not Found"

        while True:
            cmd, data = self.cmd_queue.get()
            if cmd == "load_page":
                try:
                    if not playwright:
                        logging.info("Initializing Playwright...")
                        self.after(0, self.update_status, "Initializing browser...", "working")
                        playwright = sync_playwright().start()
                        browser = playwright.chromium.launch(headless=True)
                        context = browser.new_context()
                        page = context.new_page()

                    test_case_run_id, user_name, password = data

                    base_url = "https://sycamoreinformatics.tpondemand.com/"
                    login_url = f"{base_url}login.aspx"
                    target_url = f"{base_url}RestUI/Board.aspx#page=testcaserun/{test_case_run_id}"

                    logging.info(f"Navigating to {target_url}")
                    self.after(0, self.update_status, "Navigating to test case run page...", "working")
                    page.goto(target_url)
                    page.wait_for_load_state("domcontentloaded", timeout=10000)

                    # Check if we are on the login page by looking for the 'Email' textbox
                    if page.is_visible('input[name*="Login$Email"]') or page.url.startswith(login_url):
                        logging.info("Login form detected. Logging in...")
                        self.after(0, self.update_status, "Logging in... Please wait.", "working")
                        page.get_by_role("textbox", name="Email").fill(user_name)
                        page.get_by_role("textbox", name="Password").fill(password)
                        page.get_by_role("button", name="Log in").click()
                        page.wait_for_url(f"**/*testcaserun/{test_case_run_id}", timeout=20000)
                    else:
                        logging.info("Already logged in.")

                    logging.info("Successfully on the Test Case Run page.")
                    self.after(0, self.update_status, "Loading parameters...", "working")

                    # Wait for dynamic content to load
                    page.wait_for_timeout(3000)

                    # Automatically extract parameters after successful login
                    try:
                        logging.info("--- Extracting Test Case ID, Test Plan Run ID, and Test Case Run Name ---")

                        # Extract Test Case Run Name from page title or heading
                        logging.info("Extracting Test Case Run Name...")
                        try:
                            # Use JavaScript to find the test case run name more reliably
                            js_code = """
                            () => {
                                // Look for the entity name in common locations
                                // Try the main content area with entity-name class
                                let nameElem = document.querySelector('.entity-name');
                                if (nameElem && nameElem.textContent.trim()) {
                                    return nameElem.textContent.trim();
                                }

                                // Try to find h1 or h2 that contains test case info
                                const headings = document.querySelectorAll('h1, h2, .tau-board-header__entity-name');
                                for (let h of headings) {
                                    const text = h.textContent.trim();
                                    // Look for text that contains underscores (typical test case name format)
                                    if (text && text.includes('_') && text.length > 5 && text.length < 200) {
                                        return text;
                                    }
                                }

                                // Try to find span or div with the test case run name
                                const spans = document.querySelectorAll('span, div');
                                for (let elem of spans) {
                                    const text = elem.textContent.trim();
                                    // Look for TC prefix or underscore pattern
                                    if (text && (text.startsWith('TC') || text.match(/^\d+_/)) &&
                                        text.includes('_') && text.length > 10 && text.length < 200) {
                                        // Make sure it's not deeply nested (likely a display name)
                                        let parent = elem.parentElement;
                                        let depth = 0;
                                        while (parent && depth < 5) {
                                            if (parent.classList.contains('tau-board-header') ||
                                                parent.classList.contains('entity-header') ||
                                                parent.tagName === 'HEADER') {
                                                return text;
                                            }
                                            parent = parent.parentElement;
                                            depth++;
                                        }
                                    }
                                }

                                // Last resort: look in page title or any prominent text
                                const title = document.title;
                                if (title && title.includes('_')) {
                                    // Extract the meaningful part from title
                                    const match = title.match(/TC\d+[^|]*/);
                                    if (match) return match[0].trim();
                                }

                                return null;
                            }
                            """

                            test_case_run_name = page.evaluate(js_code)

                            if test_case_run_name:
                                # Clean up the name - remove ID numbers if present at the start
                                test_case_run_name = re.sub(r'^#?\d+\s*[-:]?\s*', '', test_case_run_name)
                                logging.info(f"✓ Test Case Run Name (raw): {test_case_run_name}")
                                # Format the name for display (remove prefix before first _, replace _ with spaces)
                                formatted_name = self._format_test_case_run_name(test_case_run_name)
                                self.after(0, self._update_entry, self.test_case_run_name_entry, formatted_name)
                                logging.info(f"✓ Test Case Run Name (formatted): {formatted_name}")
                            else:
                                logging.warning("Could not find Test Case Run Name on page")
                                self.after(0, self._update_entry, self.test_case_run_name_entry, "")
                        except Exception as e:
                            logging.warning(f"Could not extract Test Case Run Name: {e}")
                            self.after(0, self._update_entry, self.test_case_run_name_entry, "")

                        # Extract Test Case ID (without clicking - just read the href)
                        logging.info("Extracting Test Case ID...")
                        test_case_link = page.locator('a[href*="#page=testcase/"]').first
                        test_case_href = test_case_link.get_attribute('href')
                        test_case_id = re.search(r'testcase/(\d+)', test_case_href).group(1)
                        self.after(0, self._update_entry, self.test_case_id_entry, test_case_id)
                        logging.info(f"✓ Test Case ID: {test_case_id}")

                        # Extract Test Plan Run ID (without clicking - just read the href)
                        logging.info("Extracting Test Plan Run ID...")
                        test_plan_run_link = page.locator('a[href*="#page=testplanrun/"]').first
                        test_plan_run_href = test_plan_run_link.get_attribute('href')
                        test_plan_run_id = re.search(r'testplanrun/(\d+)', test_plan_run_href).group(1)
                        self.after(0, self._update_entry, self.test_plan_run_id_entry, test_plan_run_id)
                        logging.info(f"✓ Test Plan Run ID: {test_plan_run_id}")

                        logging.info("✓✓✓ IDs loaded successfully!")

                        # Automatically load session parameters
                        logging.info("--- Starting Export Parameters Extraction ---")
                        self.after(0, self.update_status, "Extracting session parameters... Please wait, this may take a moment.", "working")
                        extracted_data = {}

                        # 1. Navigate to TestCaseRun page and extract data
                        test_case_run_url = f"https://sycamoreinformatics.tpondemand.com/RestUI/Board.aspx#page=testcaserun/{test_case_run_id}"
                        logging.info(f"Step 1: Navigating to TestCaseRun page...")
                        self.after(0, self.update_status, "Loading Test Case Run page... (1/3)", "working")
                        page.goto(test_case_run_url)
                        page.wait_for_load_state("domcontentloaded", timeout=20000)
                        page.wait_for_timeout(3000)  # Wait for dynamic content
                        logging.info("Extracting Export Parameters from TestCaseRun Page...")
                        for ui_label, web_label in self.field_mapping.items():
                            if web_label in ["Tester", "Run Result"]:
                                extracted_data[ui_label] = get_field_value(page, web_label)
                        for ui_label, web_label in self.date_field_mapping.items():
                            if web_label in ["Start Run Date", "End Run Date"]:
                                extracted_data[ui_label] = get_field_value(page, web_label)

                        # 2. Navigate to TestCase page and extract data
                        test_case_url = f"https://sycamoreinformatics.tpondemand.com/RestUI/Board.aspx#page=testcase/{test_case_id}"
                        logging.info(f"Step 2: Navigating to Test Case page...")
                        self.after(0, self.update_status, "Loading Test Case page... (2/3)", "working")
                        page.goto(test_case_url)
                        page.wait_for_load_state("domcontentloaded", timeout=20000)
                        page.wait_for_timeout(3000)  # Wait for dynamic content
                        logging.info("Extracting Export Parameters from TestCase Page...")
                        for ui_label, web_label in self.field_mapping.items():
                            if web_label in ["Creator", "Modified By", "Services Reviewer", "ProdMgmt Reviewer", "Head of Eng. Reviewer"]:
                                extracted_data[ui_label] = get_field_value(page, web_label)
                        for ui_label, web_label in self.date_field_mapping.items():
                            if web_label in ["Creation date", "Last Modified"]:
                                extracted_data[ui_label] = get_field_value(page, web_label)

                        # 3. Navigate to TestPlanRun page and extract data
                        test_plan_run_url = f"https://sycamoreinformatics.tpondemand.com/RestUI/Board.aspx#page=testplanrun/{test_plan_run_id}"
                        logging.info(f"Step 3: Navigating to Test Plan Run page...")
                        self.after(0, self.update_status, "Loading Test Plan Run page... (3/3)", "working")
                        page.goto(test_plan_run_url)
                        page.wait_for_load_state("domcontentloaded", timeout=20000)
                        page.wait_for_timeout(3000)  # Wait for dynamic content
                        logging.info("Extracting Export Parameters from TestPlanRun Page...")
                        for ui_label, web_label in self.field_mapping.items():
                            if web_label in ["Test Machine OS", "Browser Name", "URL"]:
                                extracted_data[ui_label] = get_field_value(page, web_label)

                        # Extract Release field using tau-linkentity__inner class
                        try:
                            release_js = """
                            () => {
                                // Find all elements containing the "Release" label
                                const allElements = Array.from(document.querySelectorAll('*'));
                                for (let elem of allElements) {
                                    if (elem.textContent && elem.textContent.trim() === 'Release') {
                                        // Found the label, now look for tau-linkentity__inner in siblings or nearby elements
                                        let parent = elem.parentElement;
                                        if (parent) {
                                            // Look in the parent's children for tau-linkentity__inner
                                            let linkEntity = parent.querySelector('.tau-linkentity__inner');
                                            if (linkEntity && linkEntity.textContent.trim()) {
                                                return linkEntity.textContent.trim();
                                            }
                                            // Also check parent's next sibling
                                            if (parent.nextElementSibling) {
                                                linkEntity = parent.nextElementSibling.querySelector('.tau-linkentity__inner');
                                                if (linkEntity && linkEntity.textContent.trim()) {
                                                    return linkEntity.textContent.trim();
                                                }
                                            }
                                        }
                                        // Check siblings
                                        if (elem.nextElementSibling) {
                                            let linkEntity = elem.nextElementSibling.querySelector('.tau-linkentity__inner');
                                            if (linkEntity && linkEntity.textContent.trim()) {
                                                return linkEntity.textContent.trim();
                                            }
                                            // Or the sibling itself might have the class
                                            if (elem.nextElementSibling.classList.contains('tau-linkentity__inner')) {
                                                return elem.nextElementSibling.textContent.trim();
                                            }
                                        }
                                    }
                                }
                                return null;
                            }
                            """
                            release_value = page.evaluate(release_js)
                            if release_value:
                                extracted_data["Release"] = release_value
                                logging.info(f"Extracted Release value: '{release_value}'")
                            else:
                                extracted_data["Release"] = "Not Found"
                                logging.warning("Could not find Release value")
                        except Exception as e:
                            logging.error(f"Failed to extract Release: {e}")
                            extracted_data["Release"] = "Not Found"

                        logging.info(f"✓✓✓ All Export Parameters loaded successfully! Data: {extracted_data}")

                        # 4. Extract test case steps (Table 5 data)
                        logging.info("Step 4: Extracting test case steps...")
                        self.after(0, self.update_status, "Extracting test case steps...", "working")
                        try:
                            # Navigate back to test case run page
                            page.goto(test_case_run_url)
                            page.wait_for_load_state("domcontentloaded", timeout=20000)
                            page.wait_for_timeout(3000)

                            # Extract Pre-Requisites from Test Case Run page
                            try:
                                prerequisites_js = """
                                () => {
                                    // Find all divs with the specific class combination (without requiring data-placeholder)
                                    const divs = document.querySelectorAll('div.ui-description__inner.tau-clientinput.i-role-property');

                                    // Debug: log what we find
                                    console.log('Found ' + divs.length + ' matching divs');

                                    // The first matching div should be the Pre-Requisites section
                                    if (divs.length > 0) {
                                        const prereqDiv = divs[0];
                                        // Get all text content, preserving structure
                                        const text = prereqDiv.textContent.trim();
                                        if (text) {
                                            return text;
                                        }
                                    }

                                    return '';
                                }
                                """
                                prerequisites_text = page.evaluate(prerequisites_js)
                                self.prerequisites_text = prerequisites_text if prerequisites_text else ""

                                # Format prerequisites: add newlines before numbered items and notes
                                if self.prerequisites_text:
                                    # Remove "Pre-Requisites" or "Pre-Requisite" header text from the beginning
                                    formatted = re.sub(r'^Pre-?Requisites?\s*', '', self.prerequisites_text, flags=re.IGNORECASE)
                                    # Add newline before numbered items (1., 2., 3., etc.) but not at the start
                                    formatted = re.sub(r'(?<!^)(\d+\.)', r'\n\1', formatted)
                                    # Add blank line before and after NOTE: or Note: (case insensitive)
                                    formatted = re.sub(r'(?i)(?<!\n)(NOTE:)', r'\n\n\1', formatted)
                                    # Add blank line after the NOTE section (before next numbered item or end)
                                    formatted = re.sub(r'(?i)(NOTE:[^\n]*?)(\n)(\d+\.)', r'\1\n\n\3', formatted)
                                    # Add blank line at the end
                                    formatted = formatted.strip() + '\n'
                                    self.prerequisites_text = formatted
                                    logging.info(f"Extracted and formatted prerequisites: '{self.prerequisites_text[:100]}...'")
                                else:
                                    logging.warning("No prerequisites found on Test Case Run page")
                            except Exception as e:
                                logging.error(f"Failed to extract prerequisites: {e}")
                                self.prerequisites_text = ""

                            # Extract test case steps using specific selectors
                            try:
                                steps_js = """
                                () => {
                                    const steps = [];
                                    // Find all test-case__input elements for actions and expected results
                                    const actionDivs = document.querySelectorAll('div.test-case__input.i-role-editable');
                                    // Find all qa-button__text spans for status
                                    const statusSpans = document.querySelectorAll('span.qa-button__text');

                                    // Iterate through actions (assuming actions and expected results alternate)
                                    for (let i = 0; i < actionDivs.length; i += 2) {
                                        const actionDiv = actionDivs[i];
                                        const expectedDiv = actionDivs[i + 1];
                                        const statusSpan = statusSpans[Math.floor(i / 2)];

                                        const actionText = actionDiv ? actionDiv.textContent.trim() : '';
                                        const expectedText = expectedDiv ? expectedDiv.textContent.trim() : '';
                                        const statusText = statusSpan ? statusSpan.textContent.trim() : '';

                                        if (actionText || expectedText) {
                                            steps.push({
                                                action: actionText,
                                                expectedResult: expectedText,
                                                status: statusText
                                            });
                                        }
                                    }
                                    return steps;
                                }
                                """
                                steps_data = page.evaluate(steps_js)
                                self.test_case_steps = steps_data if steps_data else []
                                logging.info(f"Extracted {len(self.test_case_steps)} test case steps")
                            except Exception as e:
                                logging.error(f"Failed to extract test case steps: {e}")
                                self.test_case_steps = []

                            # Extract screenshots/additional info from tiptap-editor-root-richeditor_395
                            try:
                                # Wait longer and try to interact with the page to load dynamic content
                                page.wait_for_timeout(2000)

                                # Try to click on tabs/sections that might contain the screenshots
                                try:
                                    # Click on additional info section
                                    page.locator("#entityadditionalinfo523").click(timeout=3000)
                                    page.wait_for_timeout(1000)
                                except:
                                    logging.info("Could not click on entityadditionalinfo523")

                                # Get all children under //*[@id="tiptap-editor-root-richeditor_395"]/div[1]/div
                                screenshots_js = """
                                () => {
                                    let result = {html: "", text: "", children: [], images: [], contentItems: [], attempts: []};

                                    // Helper function to extract ordered content (text and images) from an element
                                    function extractOrderedContent(element) {
                                        const contentItems = [];

                                        // Walk through all child nodes in order
                                        function processNode(node, currentText = '') {
                                            if (node.nodeType === Node.TEXT_NODE) {
                                                const text = node.textContent.trim();
                                                if (text) {
                                                    return currentText + text + ' ';
                                                }
                                            } else if (node.nodeType === Node.ELEMENT_NODE) {
                                                if (node.tagName === 'IMG') {
                                                    // Save any accumulated text before the image
                                                    if (currentText.trim()) {
                                                        contentItems.push({
                                                            type: 'text',
                                                            content: currentText.trim()
                                                        });
                                                        currentText = '';
                                                    }

                                                    // Add the image
                                                    contentItems.push({
                                                        type: 'image',
                                                        src: node.src,
                                                        alt: node.alt || '',
                                                        width: node.width || node.naturalWidth,
                                                        height: node.height || node.naturalHeight,
                                                        style: node.getAttribute('style') || ''
                                                    });
                                                } else if (node.tagName === 'BR') {
                                                    currentText += '\\n';
                                                } else {
                                                    // Process children of this element
                                                    for (let child of node.childNodes) {
                                                        currentText = processNode(child, currentText);
                                                    }

                                                    // Add spacing after block elements
                                                    if (['P', 'DIV', 'H1', 'H2', 'H3', 'H4', 'H5', 'H6'].includes(node.tagName)) {
                                                        if (currentText.trim()) {
                                                            contentItems.push({
                                                                type: 'text',
                                                                content: currentText.trim()
                                                            });
                                                            currentText = '';
                                                        }
                                                    }
                                                }
                                            }
                                            return currentText;
                                        }

                                        // Process all children
                                        let remainingText = '';
                                        for (let child of element.childNodes) {
                                            remainingText = processNode(child, remainingText);
                                        }

                                        // Add any remaining text
                                        if (remainingText.trim()) {
                                            contentItems.push({
                                                type: 'text',
                                                content: remainingText.trim()
                                            });
                                        }

                                        return contentItems;
                                    }

                                    // Helper function to extract images (for backward compatibility)
                                    function extractImages(element) {
                                        const images = [];
                                        const imgTags = element.querySelectorAll('img');
                                        for (let img of imgTags) {
                                            images.push({
                                                src: img.src,
                                                alt: img.alt || '',
                                                width: img.width || img.naturalWidth,
                                                height: img.height || img.naturalHeight,
                                                style: img.getAttribute('style') || ''
                                            });
                                        }
                                        return images;
                                    }

                                    // Approach 1: Specific path as requested
                                    const editorRoot = document.querySelector('#tiptap-editor-root-richeditor_395');
                                    result.attempts.push({method: 'editorRoot', found: !!editorRoot});

                                    if (editorRoot) {
                                        const firstDiv = editorRoot.querySelector('div:first-child');
                                        result.attempts.push({method: 'firstDiv', found: !!firstDiv});

                                        if (firstDiv) {
                                            const targetDiv = firstDiv.querySelector('div');
                                            result.attempts.push({method: 'targetDiv', found: !!targetDiv});

                                            if (targetDiv && (targetDiv.innerHTML || targetDiv.textContent)) {
                                                const children = [];
                                                for (let child of targetDiv.children) {
                                                    children.push({
                                                        tagName: child.tagName,
                                                        innerHTML: child.innerHTML,
                                                        textContent: child.textContent
                                                    });
                                                }

                                                result.html = targetDiv.innerHTML;
                                                result.text = targetDiv.textContent || targetDiv.innerText;
                                                result.children = children;
                                                result.images = extractImages(targetDiv);
                                                result.contentItems = extractOrderedContent(targetDiv);
                                                result.foundMethod = 'specific-path';
                                                return result;
                                            }
                                        }
                                    }

                                    // Approach 2: Try any richeditor
                                    const allEditors = document.querySelectorAll('[id*="richeditor"]');
                                    result.attempts.push({method: 'allEditors', count: allEditors.length});
                                    for (let editor of allEditors) {
                                        if (editor.textContent && editor.textContent.trim().length > 20) {
                                            result.html = editor.innerHTML;
                                            result.text = editor.textContent;
                                            result.images = extractImages(editor);
                                            result.contentItems = extractOrderedContent(editor);
                                            result.foundMethod = 'richeditor-' + editor.id;
                                            return result;
                                        }
                                    }

                                    // Approach 3: Try tiptap editors
                                    const tiptapEditors = document.querySelectorAll('[id*="tiptap"]');
                                    result.attempts.push({method: 'tiptapEditors', count: tiptapEditors.length});
                                    for (let editor of tiptapEditors) {
                                        if (editor.textContent && editor.textContent.trim().length > 20) {
                                            result.html = editor.innerHTML;
                                            result.text = editor.textContent;
                                            result.images = extractImages(editor);
                                            result.contentItems = extractOrderedContent(editor);
                                            result.foundMethod = 'tiptap-' + editor.id;
                                            return result;
                                        }
                                    }

                                    return result;
                                }
                                """
                                screenshots_data = page.evaluate(screenshots_js)

                                # Log attempts
                                logging.info(f"Screenshot extraction attempts: {screenshots_data.get('attempts', [])}")

                                # Store the content
                                if screenshots_data.get('html') and screenshots_data.get('html').strip():
                                    self.screenshots_text = screenshots_data.get('html', '')
                                    # Clean the text: remove "View Full Size " and "&nbsp;"
                                    self.screenshots_text = self.screenshots_text.replace('View Full Size ', '').replace('View Full Size', '')
                                    self.screenshots_text = self.screenshots_text.replace('&nbsp;', ' ').replace('\xa0', ' ')
                                    logging.info(f"✓ Found screenshots via {screenshots_data.get('foundMethod', 'unknown')}: {len(self.screenshots_text)} chars")
                                elif screenshots_data.get('text') and screenshots_data.get('text').strip():
                                    self.screenshots_text = screenshots_data.get('text', '')
                                    # Clean the text: remove "View Full Size " and "&nbsp;"
                                    self.screenshots_text = self.screenshots_text.replace('View Full Size ', '').replace('View Full Size', '')
                                    self.screenshots_text = self.screenshots_text.replace('&nbsp;', ' ').replace('\xa0', ' ')
                                    logging.info(f"✓ Found text via {screenshots_data.get('foundMethod', 'unknown')}: {len(self.screenshots_text)} chars")
                                else:
                                    self.screenshots_text = ""
                                    logging.warning(f"✗ No screenshots found. Attempts: {screenshots_data.get('attempts', [])}")

                                # Store images separately
                                images_data = screenshots_data.get('images', [])
                                if images_data:
                                    logging.info(f"✓ Found {len(images_data)} images in screenshots")
                                    self.screenshots_images = []
                                    for img in images_data:
                                        img_src = img.get('src', '')
                                        if img_src:
                                            # Download or process image
                                            try:
                                                import base64

                                                # Clean caption text
                                                caption = img.get('caption', '')
                                                if caption:
                                                    # Remove "View Full Size " and "&nbsp;"
                                                    caption = caption.replace('View Full Size ', '').replace('View Full Size', '')
                                                    caption = caption.replace('&nbsp;', ' ').replace('\xa0', ' ')
                                                    caption = ' '.join(caption.split())  # Normalize whitespace
                                                    caption = caption.strip()

                                                # Check if it's a base64 image
                                                if img_src.startswith('data:image'):
                                                    # Extract base64 data
                                                    _, data = img_src.split(',', 1)
                                                    image_bytes = base64.b64decode(data)
                                                    self.screenshots_images.append({
                                                        'type': 'base64',
                                                        'data': image_bytes,
                                                        'alt': img.get('alt', ''),
                                                        'width': img.get('width'),
                                                        'height': img.get('height'),
                                                        'caption': caption
                                                    })
                                                    logging.info(f"  - Extracted base64 image ({len(image_bytes)} bytes)")
                                                elif img_src.startswith('http'):
                                                    # It's a URL - download the image
                                                    try:
                                                        logging.info(f"  - Downloading image from URL: {img_src[:100]}")
                                                        # Use Playwright to download the image
                                                        response = page.request.get(img_src)
                                                        if response.ok:
                                                            image_bytes = response.body()
                                                            self.screenshots_images.append({
                                                                'type': 'base64',  # Treat downloaded images as base64/binary
                                                                'data': image_bytes,
                                                                'alt': img.get('alt', ''),
                                                                'width': img.get('width'),
                                                                'height': img.get('height'),
                                                                'caption': caption
                                                            })
                                                            if caption:
                                                                logging.info(f"  - ✓ Downloaded image ({len(image_bytes)} bytes) with caption: {caption[:50]}")
                                                            else:
                                                                logging.info(f"  - ✓ Downloaded image ({len(image_bytes)} bytes)")
                                                        else:
                                                            logging.warning(f"  - ✗ Failed to download image (status {response.status})")
                                                    except Exception as download_error:
                                                        logging.error(f"  - ✗ Error downloading image: {download_error}")
                                            except Exception as img_error:
                                                logging.error(f"Failed to process image: {img_error}")
                                else:
                                    self.screenshots_images = []
                                    logging.info("No images found in screenshots")

                                # Process ordered content items (text and images in sequence)
                                content_items_data = screenshots_data.get('contentItems', [])
                                if content_items_data:
                                    logging.info(f"✓ Found {len(content_items_data)} content items (text and images)")
                                    self.screenshots_content_items = []

                                    for idx, item in enumerate(content_items_data):
                                        try:
                                            if item.get('type') == 'text':
                                                # Store text content
                                                text_content = item.get('content', '')
                                                # Clean text
                                                text_content = text_content.replace('View Full Size ', '').replace('View Full Size', '')
                                                text_content = text_content.replace('&nbsp;', ' ').replace('\xa0', ' ')
                                                self.screenshots_content_items.append({
                                                    'type': 'text',
                                                    'content': text_content
                                                })
                                            elif item.get('type') == 'image':
                                                # Download and store image
                                                img_src = item.get('src', '')
                                                if img_src:
                                                    if img_src.startswith('data:image'):
                                                        # Base64 image
                                                        _, data = img_src.split(',', 1)
                                                        image_bytes = base64.b64decode(data)
                                                        self.screenshots_content_items.append({
                                                            'type': 'image',
                                                            'data': image_bytes,
                                                            'width': item.get('width'),
                                                            'height': item.get('height')
                                                        })
                                                    elif img_src.startswith('http'):
                                                        # Download image
                                                        response = page.request.get(img_src)
                                                        if response.ok:
                                                            image_bytes = response.body()
                                                            self.screenshots_content_items.append({
                                                                'type': 'image',
                                                                'data': image_bytes,
                                                                'width': item.get('width'),
                                                                'height': item.get('height')
                                                            })
                                        except Exception as item_error:
                                            logging.error(f"Failed to process content item {idx}: {item_error}")

                                    logging.info(f"✓ Processed {len(self.screenshots_content_items)} ordered content items")
                                else:
                                    self.screenshots_content_items = []
                                    logging.info("No ordered content items found")

                                if self.screenshots_text:
                                    preview = self.screenshots_text[:300].replace('\n', ' ')
                                    logging.info(f"Preview: {preview}...")
                            except Exception as e:
                                logging.error(f"Failed to extract screenshots: {e}", exc_info=True)
                                self.screenshots_text = ""

                        except Exception as e:
                            logging.error(f"Failed to extract test case steps: {e}")
                            self.test_case_steps = []
                            self.screenshots_text = ""

                        # Check if any required fields are missing
                        missing_fields = [k for k, v in extracted_data.items() if v == "Not Found"]
                        if missing_fields:
                            warning_msg = f"⚠ Warning: Could not find {len(missing_fields)} field(s): {', '.join(missing_fields[:3])}"
                            if len(missing_fields) > 3:
                                warning_msg += f" and {len(missing_fields)-3} more"
                            self.after(0, self.update_status, warning_msg, "warning")
                        else:
                            self.after(0, self.update_status, "✓ All parameters loaded successfully! You can now Preview or Export.", "success")

                        self.after(0, self._update_ui_with_data, extracted_data)

                        # Always enable export buttons after parameters are loaded
                        images_with_data = [img for img in self.screenshots_images if img.get('data') and len(img.get('data', b'')) > 0]

                        logging.info(f"===== EXTRACTION COMPLETE =====")
                        if images_with_data:
                            logging.info(f"✓ Screenshots available: {len(images_with_data)} images with data extracted (out of {len(self.screenshots_images)} total)")
                            for idx, img in enumerate(images_with_data[:5]):  # Log first 5 images
                                img_size = len(img.get('data', b''))
                                logging.info(f"  - Image {idx + 1}: {img_size} bytes, type={img.get('type')}")
                            if len(images_with_data) > 5:
                                logging.info(f"  ... and {len(images_with_data) - 5} more images")
                        else:
                            if self.screenshots_images:
                                logging.info(f"⚠ Found {len(self.screenshots_images)} images but none have data.")
                            else:
                                logging.info("ℹ No screenshots found in this test case.")

                        # Always enable export buttons
                        logging.info(f"✓ Enabling export buttons")
                        self.after(0, self._enable_export_buttons)

                    except Exception as e:
                        logging.error("!!! FAILED to load parameters !!!", exc_info=True)
                        error_msg = "✗ Failed to extract parameters. Please check if the page loaded correctly."
                        self.after(0, self.update_status, error_msg, "error")

                except Exception as e:
                    logging.error(f"Failed to load page: {e}", exc_info=True)
                    self.after(0, self.update_status, f"✗ Failed to load page: {str(e)}", "error")

            elif cmd == "stop":
                logging.info("Stopping Playwright worker.")
                if browser:
                    browser.close()
                if playwright:
                    playwright.stop()
                break

    def load_page(self):
        test_case_run_id = self.id_entry.get()
        user_name = self.user_name_entry.get()
        password = self.password_entry.get()

        if not test_case_run_id or not user_name or not password:
            logging.error("Test Case Run ID, User Name, and Password are required.")
            self.update_status("Error: Please fill in Test Case Run ID, User Name, and Password", "error")
            return

        self.update_status("Loading page and logging in... Please wait, do not close the app.", "working")
        self.cmd_queue.put(("load_page", (test_case_run_id, user_name, password)))

    def _update_entry(self, entry, value):
        entry.delete(0, "end")
        entry.insert(0, value)

    def _format_script_name(self, test_case_run_name):
        """Format the script name from Test Script Name
        The test_case_run_name is already formatted with RTVE prefix in UI, so return as-is
        """
        if not test_case_run_name:
            return "RTVE - [No Name]"

        # The name is already formatted with "RTVE - " prefix, return as-is
        return test_case_run_name

    def _format_test_case_run_name(self, raw_name):
        """Format the Test Script Name for display
        - Remove everything before the first underscore
        - Replace underscores with spaces
        - Add "RTVE - " prefix
        """
        if not raw_name:
            return ""

        # Remove everything before first underscore (including the underscore)
        if '_' in raw_name:
            formatted_name = raw_name.split('_', 1)[1]
        else:
            formatted_name = raw_name

        # Replace underscores with spaces
        formatted_name = formatted_name.replace('_', ' ')

        # Add "RTVE - " prefix
        formatted_name = "RTVE - " + formatted_name

        return formatted_name

    def _clean_extracted_value(self, ui_label, value):
        """Remove label prefix from extracted value if present"""
        # Get the web label for this UI field
        web_label = None

        # Check in field_mapping (UI label -> web label)
        if ui_label in self.field_mapping:
            web_label = self.field_mapping[ui_label]
        # Check in date_field_mapping
        elif ui_label in self.date_field_mapping:
            web_label = self.date_field_mapping[ui_label]

        # If we found the web label, try to remove it from the value
        if web_label and value.startswith(web_label):
            # Remove the label and any spaces/tabs after it
            cleaned = value[len(web_label):].strip()
            logging.info(f"Cleaned '{ui_label}': '{value}' -> '{cleaned}'")
            return cleaned

        return value

    def _update_ui_with_data(self, data):
        logging.info("Updating UI with extracted data.")
        for field, value in data.items():
            if value == "Not Found":
                logging.warning(f"No value found for '{field}'. Skipping UI update.")
                continue

            # Clean the value to remove label prefixes
            value = self._clean_extracted_value(field, value)

            if field in self.form_entries:
                self._update_entry(self.form_entries[field], value)
            elif field == "Release":
                # Extract only the version number (e.g., "4.5.0.0" from "Sycamore CDR-SCE 4.5.0.0")
                version_match = re.search(r'(\d+(?:\.\d+)+)', value)
                if version_match:
                    value = version_match.group(1)
                self._update_entry(self.release_entry, value)
            elif field in self.date_entries:
                try:
                    # Attempt to parse dates like "01-Dec-2025" or "1-Dec-2025 07:56"
                    match = re.search(r"(\d{1,2})-([A-Za-z]{3})-(\d{4})", value)
                    if match:
                        date_str = match.group(0)
                        dt_object = datetime.strptime(date_str, '%d-%b-%Y')
                        self.date_entries[field].set_date(dt_object)
                        logging.info(f"Set date for '{field}' to {dt_object.strftime('%Y-%m-%d')}")
                    else:
                        logging.warning(f"Could not parse date string for '{field}': '{value}'")
                except Exception as e:
                    logging.error(f"Error setting date for '{field}': {e}", exc_info=True)
            else:
                logging.warning(f"UI element for field '{field}' not found.")

    def _enable_export_buttons(self):
        """Enable the export buttons when screenshots are available"""
        self.export_word_button.configure(state="normal")
        self.export_pdf_button.configure(state="normal")
        logging.info("Export buttons enabled")

    def _disable_export_buttons(self):
        """Disable the export buttons when screenshots are not available"""
        self.export_word_button.configure(state="disabled")
        self.export_pdf_button.configure(state="disabled")
        logging.info("Export buttons disabled")

    def preview_data(self):
        """Preview all collected data"""
        logging.info("Preview button clicked")

        # Validate that Test Case Run Name is available
        test_case_run_name = self.test_case_run_name_entry.get()
        if not test_case_run_name:
            self.update_status("✗ Error: Test Script Name not loaded. Please load parameters first.", "error")
            return

        self.update_status("Generating preview...", "working")

        try:
            # Get current year
            from datetime import datetime
            current_year = datetime.now().year

            # Format script name
            script_name = self._format_script_name(test_case_run_name)

            # Collect all form data
            data = self._collect_form_data()

            # Generate preview text
            preview = self._generate_preview_text(script_name, current_year, data)

            # Display in preview area
            self.preview_text.delete("1.0", "end")
            self.preview_text.insert("1.0", preview)

            logging.info("Preview generated successfully")
            self.update_status("✓ Preview generated successfully!", "success")

        except Exception as e:
            logging.error(f"Failed to generate preview: {e}", exc_info=True)
            self.update_status(f"✗ Failed to generate preview: {str(e)}", "error")

    def _collect_form_data(self):
        """Collect all data from form fields"""
        data = {}

        # Collect text entries
        for label, entry in self.form_entries.items():
            data[label] = entry.get() if entry.get() else "[Not Set]"

        # Collect date entries
        for label, date_entry in self.date_entries.items():
            try:
                date_obj = date_entry.get_date()
                data[label] = date_obj.strftime('%d-%b-%Y')
            except:
                data[label] = "[Not Set]"

        # Collect Release field
        data['Release'] = self.release_entry.get() if self.release_entry.get() else "[Not Set]"

        return data

    def _generate_preview_text(self, script_name, current_year, data):
        """Generate formatted preview text"""
        preview = []
        W = 120  # Total width

        # Header
        preview.append("=" * W)
        header_left = "Sycamore Informatics"
        spaces = W - len(header_left) - len(script_name) - 2
        preview.append(f"  {header_left}{' ' * spaces}{script_name}")
        preview.append("=" * W)
        preview.append("")
        preview.append(" " * 55 + "PAGE 2")
        preview.append("")
        preview.append("")

        # Table 1: Author Details (5 rows × 2 columns)
        preview.append("╔" + "═" * (W-2) + "╗")
        title = "AUTHOR DETAILS"
        preview.append("║" + " " * ((W-2-len(title))//2) + title + " " * ((W-2-len(title)+1)//2) + "║")
        preview.append("║" + " " * ((W-2-len("[BLUE HEADER]"))//2) + "[BLUE HEADER]" + " " * ((W-2-len("[BLUE HEADER]")+1)//2) + "║")
        preview.append("╠" + "═" * ((W-2)//2) + "╦" + "═" * ((W-2)//2) + "╣")

        # Row 2: Author Name
        label = " Author Name: [GRAY BG]"
        value = " " + data.get('Author Name', '[Not Set]')
        preview.append(f"║{label:<{(W-2)//2}}║{value:<{(W-2)//2}}║")
        preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        # Row 3: Author Date
        label = " Author Date: [GRAY BG]"
        value = " " + data.get('Author Date', '[Not Set]')
        preview.append(f"║{label:<{(W-2)//2}}║{value:<{(W-2)//2}}║")
        preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        # Row 4: Last Updated By
        label = " Last Updated By: [GRAY BG]"
        value = " " + data.get('Last Updated By', '[Not Set]')
        preview.append(f"║{label:<{(W-2)//2}}║{value:<{(W-2)//2}}║")
        preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        # Row 5: Last Updated Date
        label = " Last Updated Date: [GRAY BG]"
        value = " " + data.get('Last Updated Date', '[Not Set]')
        preview.append(f"║{label:<{(W-2)//2}}║{value:<{(W-2)//2}}║")
        preview.append("╚" + "═" * ((W-2)//2) + "╩" + "═" * ((W-2)//2) + "╝")
        preview.append("")
        preview.append("")

        # Table 2: Reviewer / Approver (5 rows × 3 columns)
        col_w = (W-2)//3
        preview.append("╔" + "═" * (W-2) + "╗")
        title = "REVIEWER / APPROVER"
        preview.append("║" + " " * ((W-2-len(title))//2) + title + " " * ((W-2-len(title)+1)//2) + "║")
        preview.append("║" + " " * ((W-2-len("[BLUE HEADER]"))//2) + "[BLUE HEADER]" + " " * ((W-2-len("[BLUE HEADER]")+1)//2) + "║")
        preview.append("╠" + "═" * col_w + "╦" + "═" * col_w + "╦" + "═" * col_w + "╣")

        # Row 2: Column Headers [GRAY BG]
        preview.append(f"║{'Name [GRAY BG]':^{col_w}}║{'Role [GRAY BG]':^{col_w}}║{'Responsibility [GRAY BG]':^{col_w}}║")
        preview.append("╠" + "═" * col_w + "╬" + "═" * col_w + "╬" + "═" * col_w + "╣")

        # Row 3: Services Approver
        name = data.get('Services Approver', '[Not Set]')[:col_w-2]
        role = "Approver (Services)"
        resp = "To review and approve the"
        resp2 = "script before the execution"
        preview.append(f"║ {name:<{col_w-2}}║ {role:<{col_w-2}}║ {resp:<{col_w-2}}║")
        preview.append(f"║{' ':<{col_w}}║{' ':<{col_w}}║ {resp2:<{col_w-2}}║")
        preview.append("╠" + "═" * col_w + "╬" + "═" * col_w + "╬" + "═" * col_w + "╣")

        # Row 4: PM Approver
        name = data.get('Product Management Approver', '[Not Set]')[:col_w-2]
        role = "Approver (Product Management)"
        preview.append(f"║ {name:<{col_w-2}}║ {role:<{col_w-2}}║ {resp:<{col_w-2}}║")
        preview.append(f"║{' ':<{col_w}}║{' ':<{col_w}}║ {resp2:<{col_w-2}}║")
        preview.append("╠" + "═" * col_w + "╬" + "═" * col_w + "╬" + "═" * col_w + "╣")

        # Row 5: Dev Approver
        name = data.get('Development Approver', '[Not Set]')[:col_w-2]
        role = "Approver (Development)"
        preview.append(f"║ {name:<{col_w-2}}║ {role:<{col_w-2}}║ {resp:<{col_w-2}}║")
        preview.append(f"║{' ':<{col_w}}║{' ':<{col_w}}║ {resp2:<{col_w-2}}║")
        preview.append("╚" + "═" * col_w + "╩" + "═" * col_w + "╩" + "═" * col_w + "╝")
        preview.append("")
        preview.append("")

        # Table 3: Execution Details (12 rows × 2 columns)
        preview.append("╔" + "═" * (W-2) + "╗")
        title = "EXECUTION DETAILS"
        preview.append("║" + " " * ((W-2-len(title))//2) + title + " " * ((W-2-len(title)+1)//2) + "║")
        preview.append("║" + " " * ((W-2-len("[BLUE HEADER]"))//2) + "[BLUE HEADER]" + " " * ((W-2-len("[BLUE HEADER]")+1)//2) + "║")
        preview.append("╠" + "═" * ((W-2)//2) + "╦" + "═" * ((W-2)//2) + "╣")

        fields = [
            ("Release: [GRAY BG]", data.get('Release', '[Not Set]')),
            ("Executed By: [GRAY BG]", data.get('Executor Name', '[Not Set]')),
            ("Execution Start Date: [GRAY BG]", data.get('Execution Start Date', '[Not Set]')),
            ("Execution End Date: [GRAY BG]", data.get('Execution End Date', '[Not Set]')),
            ("Execution Status: [GRAY BG]", data.get('Execution Status', '[Not Set]')),
            ("Test Operating System: [GRAY BG]", data.get('Test Operating System', '[Not Set]')),
            ("Browser: [GRAY BG]", data.get('Browser Name', '[Not Set]')),
            ("URL: [GRAY BG]", data.get('URL', '[Not Set]'))
        ]

        for i, (label, value) in enumerate(fields):
            label_cell = " " + label
            value_cell = " " + value
            preview.append(f"║{label_cell:<{(W-2)//2}}║{value_cell:<{(W-2)//2}}║")
            if i < len(fields) - 1:
                preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        # Executor Signature (3 merged rows)
        label_cell = " Executor Signature: [GRAY BG, MERGED 3 ROWS]"
        preview.append(f"║{label_cell:<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append(f"║{' ':<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append(f"║{' ':<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append("╚" + "═" * ((W-2)//2) + "╩" + "═" * ((W-2)//2) + "╝")
        preview.append("")
        preview.append("")

        # Table 4: Services Approval Sign-Off (5 rows × 2 columns)
        preview.append("╔" + "═" * (W-2) + "╗")
        title = "SERVICES APPROVAL SIGN-OFF"
        preview.append("║" + " " * ((W-2-len(title))//2) + title + " " * ((W-2-len(title)+1)//2) + "║")
        preview.append("║" + " " * ((W-2-len("[BLUE HEADER]"))//2) + "[BLUE HEADER]" + " " * ((W-2-len("[BLUE HEADER]")+1)//2) + "║")
        preview.append("╠" + "═" * ((W-2)//2) + "╦" + "═" * ((W-2)//2) + "╣")

        # Row 2: Column Headers [GRAY BG]
        preview.append(f"║ Approver's Name [GRAY BG]{' ':<{(W-2)//2-25}}║ Signature [GRAY BG]{' ':<{(W-2)//2-19}}║")
        preview.append("╠" + "═" * ((W-2)//2) + "╬" + "═" * ((W-2)//2) + "╣")

        # Rows 3-5: Services Approver name (merged, gray BG) and signature space
        approver_name = " " + data.get('Services Approver', '[Not Set]') + " [GRAY BG, BOLD, MERGED 3 ROWS]"
        preview.append(f"║{approver_name:<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append(f"║{' ':<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append(f"║{' ':<{(W-2)//2}}║{' ':<{(W-2)//2}}║")
        preview.append("╚" + "═" * ((W-2)//2) + "╩" + "═" * ((W-2)//2) + "╝")
        preview.append("")
        preview.append("")

        # Page break marker for Page 4
        preview.append(" " * 55 + "PAGE 4")
        preview.append("")
        preview.append("")

        # Table 5: Test Case Steps (dynamic rows × 4 columns: 5%, 40%, 40%, 15%)
        c1_w = 6   # 5% - Step#
        c2_w = 48  # 40% - Action
        c3_w = 48  # 40% - Expected Result
        c4_w = 14  # 15% - Status

        preview.append("╔" + "═" * (W-2) + "╗")
        title = "PRE-REQUISITES"
        preview.append("║" + " " * ((W-2-len(title))//2) + title + " " * ((W-2-len(title)+1)//2) + "║")
        preview.append("║" + " " * ((W-2-len("[BLUE HEADER, BOLD, +2 FONT]"))//2) + "[BLUE HEADER, BOLD, +2 FONT]" + " " * ((W-2-len("[BLUE HEADER, BOLD, +2 FONT]")+1)//2) + "║")
        preview.append("╠" + "═" * (W-2) + "╣")

        # Row 2: Pre-Requisites content (merged, italic)
        prereq_text = self.prerequisites_text if self.prerequisites_text else "No prerequisites specified"
        # Wrap text to fit width while preserving intentional newlines
        import textwrap
        preview.append("║ [ITALIC, Extra line before]" + " " * (W-2-len(" [ITALIC, Extra line before]")) + "║")
        # Split by newlines first, then wrap each line individually
        for prereq_line in prereq_text.split('\n'):
            if prereq_line.strip():  # Skip empty lines in preview display
                wrapped_lines = textwrap.fill(prereq_line, width=W-6)
                for line in wrapped_lines.split('\n'):
                    display_line = line[:W-4]  # Truncate if too long
                    preview.append("║ " + display_line + " " * (W-4-len(display_line)) + "║")
        preview.append("║ [ITALIC, Extra line after]" + " " * (W-2-len(" [ITALIC, Extra line after]")) + "║")
        preview.append("╠" + "═" * c1_w + "╦" + "═" * c2_w + "╦" + "═" * c3_w + "╦" + "═" * c4_w + "╣")

        # Row 3: Column headers
        preview.append(f"║{'Step#':^{c1_w}}║{'Action':^{c2_w}}║{'Expected Result':^{c3_w}}║{'Status':^{c4_w}}║")
        preview.append(f"║{'[BLUE]':^{c1_w}}║{'[BLUE, BOLD, +1 FONT]':^{c2_w}}║{'[BLUE, BOLD, +1 FONT]':^{c3_w}}║{'[BLUE]':^{c4_w}}║")
        preview.append("╠" + "═" * c1_w + "╬" + "═" * c2_w + "╬" + "═" * c3_w + "╬" + "═" * c4_w + "╣")

        # Data rows
        step_num = 1
        if self.test_case_steps and len(self.test_case_steps) > 0:
            for step_data in self.test_case_steps[:5]:  # Show first 5 steps in preview
                action = step_data.get('action', '')[:c2_w-2]
                expected = step_data.get('expectedResult', '')[:c3_w-2]
                status = step_data.get('status', '')[:c4_w-2]

                # Check if scenario
                if 'Scenario' in action or 'SCENARIO' in action.upper():
                    # Merged row
                    preview.append(f"║{action:^{W-2}}║")
                    preview.append(f"║{'[MERGED, BOLD]':^{W-2}}║")
                else:
                    preview.append(f"║{str(step_num):^{c1_w}}║ {action:<{c2_w-2}}║ {expected:<{c3_w-2}}║{status:^{c4_w}}║")
                    step_num += 1

                preview.append("╠" + "═" * c1_w + "╬" + "═" * c2_w + "╬" + "═" * c3_w + "╬" + "═" * c4_w + "╣")

            # Indicate more rows if applicable
            if len(self.test_case_steps) > 5:
                preview.append(f"║{f'... ({len(self.test_case_steps)-5} more steps)':^{W-2}}║")
                preview.append("╠" + "═" * (W-2) + "╣")
        else:
            preview.append(f"║{'No test case steps available':^{W-2}}║")
            preview.append("╠" + "═" * (W-2) + "╣")

        # Last row: Screenshots
        screenshots_text = self.screenshots_text if self.screenshots_text else "No screenshots available"
        # Strip HTML and truncate
        import re
        clean_screenshots = re.sub('<[^<]+?>', '', screenshots_text)[:100]
        preview.append(f"║ Screenshots: {clean_screenshots:<{W-15}}║")
        preview.append(f"║{'[MERGED ROW - tiptap-editor-root-richeditor_395 content]':^{W-2}}║")
        preview.append("╚" + "═" * (W-2) + "╝")
        preview.append("")

        # Footer
        preview.append("=" * W)
        footer_text = f"© Sycamore Informatics {current_year}. Proprietary and Confidential Document, all rights reserved."
        page_num = "Page 4"
        spaces = W - len(footer_text) - len(page_num) - 4
        preview.append(f"  {footer_text}{' ' * spaces}{page_num}")
        preview.append("=" * W)

        return "\n".join(preview)

    def export_to_word(self):
        """Export data to Word document"""
        logging.info("Export to Word button clicked")

        # Log image availability
        if self.screenshots_images:
            logging.info(f"✓ {len(self.screenshots_images)} screenshot image(s) available for Word export")
            for idx, img in enumerate(self.screenshots_images):
                logging.info(f"  Image {idx + 1}: type={img.get('type')}, data_size={len(img.get('data', b''))} bytes")
        else:
            logging.warning("⚠ No screenshot images available for Word export")

        # Validate that Test Case Run Name is available
        test_case_run_name = self.test_case_run_name_entry.get()
        if not test_case_run_name:
            self.update_status("✗ Error: Test Script Name not loaded. Please load parameters first.", "error")
            return

        self.update_status("Exporting to Word... Please wait.", "working")

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import os
            from datetime import datetime

            # Get current year
            current_year = datetime.now().year

            # Format script name
            script_name = self._format_script_name(test_case_run_name)

            # Collect form data
            data = self._collect_form_data()

            # Create Word document
            doc = Document()

            # Create first page
            self._create_word_first_page(doc, script_name)

            # Set different first page for headers/footers
            section = doc.sections[0]
            section.different_first_page_header_footer = True

            # Add page break for Page 2
            doc.add_page_break()

            # Set up headers and footers for pages 2+ (add new section)
            section = doc.sections[-1]
            self._setup_word_header_footer(doc, script_name, current_year)

            # Add Table 1: Author Details
            self._add_author_details_table(doc, data)
            doc.add_paragraph()  # Blank row
            doc.add_paragraph()  # Blank row

            # Add Table 2: Reviewer / Approver Roles
            self._add_reviewer_approver_table(doc, data)

            # Add page break for Page 3
            doc.add_page_break()

            # Add Table 3: Execution Details
            self._add_execution_details_table(doc, data)
            doc.add_paragraph()  # Blank row
            doc.add_paragraph()  # Blank row

            # Add Table 4: Services Approval Sign-Off
            self._add_approval_signoff_table(doc, data)

            # Add page break for Page 4
            doc.add_page_break()

            # Verify images are still available before creating table
            if self.screenshots_images:
                logging.info(f"[Word Export] Verified: {len(self.screenshots_images)} images ready for Table 5")
            else:
                logging.warning(f"[Word Export] Warning: No images available for Table 5")

            # Add Table 5: Test Case Steps
            self._add_test_case_table(doc, data)

            # Save document with file dialog
            default_filename = f"{script_name}.docx"
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save Word Document"
            )

            if not filepath:
                self.update_status("Export cancelled by user.", "error")
                return

            # Log export summary
            images_with_data_count = len([img for img in self.screenshots_images if img.get('data') and len(img.get('data', b'')) > 0]) if self.screenshots_images else 0
            if images_with_data_count > 0:
                logging.info(f"[Word Export] ===== EXPORT COMPLETE: {images_with_data_count} images included in document =====")
            else:
                logging.info(f"[Word Export] ===== EXPORT COMPLETE: No images in document =====")

            doc.save(filepath)
            logging.info(f"Word document saved to: {filepath}")
            print(f"✓ Word document exported to: {filepath}")

            # Update status with image count
            status_msg = f"✓ Word document saved: {os.path.basename(filepath)}"
            if images_with_data_count > 0:
                status_msg += f" (with {images_with_data_count} image(s))"
            self.update_status(status_msg, "success")

        except ImportError:
            error_msg = "python-docx library not installed. Please run: pip install python-docx"
            logging.error(error_msg)
            print(f"✗ {error_msg}")
            self.update_status(f"✗ {error_msg}", "error")
        except Exception as e:
            logging.error(f"Failed to export to Word: {e}", exc_info=True)
            print(f"✗ Failed to export to Word: {e}")
            self.update_status(f"✗ Failed to export to Word: {str(e)}", "error")

    def _setup_word_header_footer(self, doc, script_name, current_year):
        """Setup header and footer for Word document"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        section = doc.sections[0]

        # Header - Create table with 2 columns (30% and 70%)
        header = section.header
        # Clear default paragraph
        header.paragraphs[0].text = ""

        # Create header table
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.autofit = False

        # Set column widths (30% and 70%)
        header_table.columns[0].width = Inches(1.95)  # 30% of 6.5 inches
        header_table.columns[1].width = Inches(4.55)  # 70% of 6.5 inches

        # Remove all table borders
        for row in header_table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tc_borders.append(border)
                tc_pr.append(tc_borders)

        # Add content to header
        left_cell = header_table.rows[0].cells[0]
        left_cell.text = "Sycamore Informatics"
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_cell.paragraphs[0].runs[0].font.size = Pt(10)

        right_cell = header_table.rows[0].cells[1]
        right_cell.text = script_name
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].runs[0].font.size = Pt(10)

        # Footer - Create table with 2 columns (90% and 10%)
        footer = section.footer
        # Clear default paragraph
        footer.paragraphs[0].text = ""

        # Create footer table
        footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        footer_table.autofit = False
        footer_table.allow_autofit = False

        # Set cell widths directly (90% and 10% of 6.5 inches)
        left_cell = footer_table.rows[0].cells[0]
        right_cell = footer_table.rows[0].cells[1]
        left_cell.width = Inches(5.85)
        right_cell.width = Inches(0.65)

        # Remove all table borders and set noWrap on cells
        for cell in [left_cell, right_cell]:
            tc_pr = cell._element.get_or_add_tcPr()
            # Set cell width explicitly in XML
            tcW = OxmlElement('w:tcW')
            if cell == left_cell:
                tcW.set(qn('w:w'), str(int(5.85 * 1440)))  # Convert inches to twips
            else:
                tcW.set(qn('w:w'), str(int(0.65 * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tc_pr.insert(0, tcW)
            # Remove borders
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tc_borders.append(border)
            tc_pr.append(tc_borders)
            # Set noWrap at cell level
            noWrap = OxmlElement('w:noWrap')
            tc_pr.append(noWrap)

        # Add content to footer - left cell (copyright)
        left_cell.text = f"© Sycamore Informatics {current_year}. Proprietary and Confidential Document, all rights reserved."
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_para.runs[0].font.size = Pt(10)

        # Right cell - page number
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add "Page " text
        run = right_para.add_run("Page ")
        run.font.size = Pt(10)

        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

    def _create_word_first_page(self, doc, script_name):
        """Create the first page with template layout"""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        import os

        # Get current date in format "22 Nov 2024"
        current_date = datetime.now().strftime("%d %b %Y")

        # Add spacing from top (2 empty paragraphs)
        doc.add_paragraph()
        doc.add_paragraph()

        # Create table for title and "By..." line (2 rows, 1 column)
        table = doc.add_table(rows=2, cols=1)
        table.style = 'Normal Table'

        # Set table width to 6.48 inches
        for row in table.rows:
            row.cells[0].width = Inches(6.48)

        # Row 0: Script name with gray background (CENTER ALIGNED)
        row0 = table.rows[0]
        row0.height = Inches(1.01)
        cell_title = row0.cells[0]

        # Clear default paragraph and add new one with proper formatting
        cell_title.text = ''
        para_title = cell_title.paragraphs[0]
        para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align script name
        para_title.paragraph_format.space_before = Pt(6)  # 0.083 inches
        para_title.paragraph_format.line_spacing = 1.05
        run_title = para_title.add_run(script_name)
        run_title.font.name = 'Trebuchet MS'
        run_title.font.size = Pt(24)

        # Set gray background for title cell
        self._set_cell_background(cell_title, "D9D9D9")

        # Row 1: "By Sycamore Informatics Inc." with blue background (RIGHT ALIGNED)
        row1 = table.rows[1]
        row1.height = Inches(0.573)
        cell_by = row1.cells[0]

        # Clear default paragraph and add new one with proper formatting
        cell_by.text = ''
        para_by = cell_by.paragraphs[0]
        para_by.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right align "By..." line
        para_by.paragraph_format.space_before = Pt(6)  # 0.083 inches
        para_by.paragraph_format.line_spacing = 1.0
        run_by = para_by.add_run("By Sycamore Informatics Inc.")
        run_by.font.name = 'Trebuchet MS'

        # Set blue background for "By" cell
        self._set_cell_background(cell_by, "9FC5E8")

        # Add spacing before logo (one empty paragraph)
        doc.add_paragraph()

        # Add logo image (CENTER ALIGNED)
        logo_path = resource_path("logo.png")
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align logo
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(4.34), height=Inches(2.20))

        # Add spacing (6 empty centered paragraphs)
        for _ in range(6):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add address on 3 separate lines (RIGHT ALIGNED, Body Text style)
        # Line 1: Address
        address_line1 = doc.add_paragraph("271 Waverley Oaks Rd #103")
        address_line1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        address_line1.style = 'Body Text'

        # Line 2: Company name
        address_line2 = doc.add_paragraph("Sycamore Informatics, Inc.")
        address_line2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        address_line2.style = 'Body Text'

        # Line 3: Date
        date_para = doc.add_paragraph(f"Date: {current_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.style = 'Body Text'

    def _add_author_details_table(self, doc, data):
        """Add Author Details table"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        # Set column widths for proper text wrapping (3" each)
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(3)

        # Row 1: Header
        cell = table.rows[0].cells[0]
        cell.merge(table.rows[0].cells[1])
        cell.text = "Author Details"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        self._set_cell_background(cell, "a0c5e8")

        # Row 2: Author Name
        self._format_label_cell(table.rows[1].cells[0], "Author Name:")
        table.rows[1].cells[1].text = data.get('Author Name', '[Not Set]')

        # Row 3: Author Date
        self._format_label_cell(table.rows[2].cells[0], "Author Date:")
        table.rows[2].cells[1].text = data.get('Author Date', '[Not Set]')

        # Row 4: Last Updated By
        self._format_label_cell(table.rows[3].cells[0], "Last Updated By:")
        table.rows[3].cells[1].text = data.get('Last Updated By', '[Not Set]')

        # Row 5: Last Updated Date
        self._format_label_cell(table.rows[4].cells[0], "Last Updated Date:")
        table.rows[4].cells[1].text = data.get('Last Updated Date', '[Not Set]')

    def _add_reviewer_approver_table(self, doc, data):
        """Add Reviewer / Approver Roles table"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        # Set column widths to ensure proper text wrapping
        # Name: 1.8", Role: 1.8", Responsibility: 2.9" (total ~6.5")
        for row in table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(1.8)
            row.cells[2].width = Inches(2.9)

        # Row 1: Header
        cell = table.rows[0].cells[0]
        cell.merge(table.rows[0].cells[1])
        cell.merge(table.rows[0].cells[2])
        cell.text = "Reviewer / Approver"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        self._set_cell_background(cell, "a0c5e8")

        # Row 2: Column headers
        headers = ["Name", "Role", "Responsibility"]
        for i, header in enumerate(headers):
            cell = table.rows[1].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(cell, "d9d9d9")

        # Row 3: Services Approver
        table.rows[2].cells[0].text = data.get('Services Approver', '[Not Set]')
        table.rows[2].cells[1].text = "Services Reviewer"
        table.rows[2].cells[2].text = "To review and approve the script before the execution"

        # Row 4: Product Management Approver
        table.rows[3].cells[0].text = data.get('Product Management Approver', '[Not Set]')
        table.rows[3].cells[1].text = "Product Management Reviewer"
        table.rows[3].cells[2].text = "To review and approve the script before the execution"

        # Row 5: Development Approver
        table.rows[4].cells[0].text = data.get('Development Approver', '[Not Set]')
        table.rows[4].cells[1].text = "Head of Engineering Reviewer"
        table.rows[4].cells[2].text = "To review and approve the script before the execution"

    def _add_execution_details_table(self, doc, data):
        """Add Execution Details table"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        table = doc.add_table(rows=12, cols=2)
        table.style = 'Table Grid'

        # Set column widths for proper text wrapping (3" each)
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(3)

        # Row 1: Header
        cell = table.rows[0].cells[0]
        cell.merge(table.rows[0].cells[1])
        cell.text = "Execution Details"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        self._set_cell_background(cell, "a0c5e8")

        # Rows 2-9: Data fields
        fields = [
            ("Release:", data.get('Release', '[Not Set]')),
            ("Executed By:", data.get('Executor Name', '[Not Set]')),
            ("Execution Start Date:", data.get('Execution Start Date', '[Not Set]')),
            ("Execution End Date:", data.get('Execution End Date', '[Not Set]')),
            ("Execution Status:", data.get('Execution Status', '[Not Set]')),
            ("Test Operating System:", data.get('Test Operating System', '[Not Set]')),
            ("Browser:", data.get('Browser Name', '[Not Set]')),
            ("URL:", data.get('URL', '[Not Set]'))
        ]

        for i, (label, value) in enumerate(fields, start=1):
            self._format_label_cell(table.rows[i].cells[0], label)
            table.rows[i].cells[1].text = value

        # Rows 10-12: Executor Signature (merged cells)
        cell_label = table.rows[9].cells[0]
        cell_label.merge(table.rows[10].cells[0])
        cell_label.merge(table.rows[11].cells[0])
        cell_label.text = "Executor Signature:"
        cell_label.paragraphs[0].runs[0].font.bold = True
        self._set_cell_background(cell_label, "d9d9d9")

        cell_value = table.rows[9].cells[1]
        cell_value.merge(table.rows[10].cells[1])
        cell_value.merge(table.rows[11].cells[1])

    def _add_approval_signoff_table(self, doc, data):
        """Add Services Approval Sign-Off table"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        # Set column widths for proper text wrapping (3" each)
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(3)

        # Row 1: Header
        cell = table.rows[0].cells[0]
        cell.merge(table.rows[0].cells[1])
        cell.text = "Services Approval Sign-Off"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        self._set_cell_background(cell, "a0c5e8")

        # Row 2: Column headers
        headers = ["Approver's Name", "Signature"]
        for i, header in enumerate(headers):
            cell = table.rows[1].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_background(cell, "d9d9d9")

        # Rows 3-5: Approver name (merged) and signature space
        cell_name = table.rows[2].cells[0]
        cell_name.merge(table.rows[3].cells[0])
        cell_name.merge(table.rows[4].cells[0])
        cell_name.text = data.get('Services Approver', '[Not Set]')
        cell_name.paragraphs[0].runs[0].font.bold = True
        self._set_cell_background(cell_name, "d9d9d9")

        cell_signature = table.rows[2].cells[1]
        cell_signature.merge(table.rows[3].cells[1])
        cell_signature.merge(table.rows[4].cells[1])

    def _add_test_case_table(self, doc, data):
        """Add Test Case Steps table (Table 5)"""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Calculate number of rows needed
        # Row 1: Pre-Requisites1 header
        # Row 2: Pre-Requisites1 content
        # Row 3: Column headers (Step#, Action, Expected Result, Status)
        # Then data rows + scenarios
        # Last row: Screenshots
        num_data_rows = len(self.test_case_steps) if self.test_case_steps else 0
        total_rows = 3 + num_data_rows + 1  # Header + prereq + column headers + data + screenshots

        # Create table with 4 columns (5%, 40%, 40%, 15%)
        table = doc.add_table(rows=total_rows, cols=4)
        table.style = 'Table Grid'

        # Set column widths: 5%, 40%, 40%, 15% (approx 0.3", 2.6", 2.6", 1")
        for row in table.rows:
            row.cells[0].width = Inches(0.3)
            row.cells[1].width = Inches(2.6)
            row.cells[2].width = Inches(2.6)
            row.cells[3].width = Inches(1.0)

        # Row 1: Pre-Requisites header (merged, bold, +2 font size, blue background)
        cell_prereq_header = table.rows[0].cells[0]
        cell_prereq_header.merge(table.rows[0].cells[1])
        cell_prereq_header.merge(table.rows[0].cells[2])
        cell_prereq_header.merge(table.rows[0].cells[3])
        cell_prereq_header.text = "Pre-Requisites"
        para = cell_prereq_header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(13)  # +2 from base 11
        self._set_cell_background(cell_prereq_header, "a0c5e8")

        # Row 2: Pre-Requisites content (merged, italic, extra lines)
        cell_prereq_content = table.rows[1].cells[0]
        cell_prereq_content.merge(table.rows[1].cells[1])
        cell_prereq_content.merge(table.rows[1].cells[2])
        cell_prereq_content.merge(table.rows[1].cells[3])

        # Add content with proper line breaks for each prerequisite
        prereq_text = self.prerequisites_text if self.prerequisites_text else "No prerequisites specified"
        # Clear the cell and add content with proper paragraphs for each line
        cell_prereq_content.text = ""  # Clear default paragraph
        prereq_lines = prereq_text.split('\n')
        for i, line in enumerate(prereq_lines):
            if i == 0:
                # First line - add to existing paragraph with extra line before
                para = cell_prereq_content.paragraphs[0]
                para.add_run("\n")  # Extra line before
                run = para.add_run(line)
                run.font.italic = True
            else:
                # Subsequent lines - add new paragraph
                para = cell_prereq_content.add_paragraph()
                run = para.add_run(line)
                run.font.italic = True
        # Add extra blank line at the end
        final_para = cell_prereq_content.add_paragraph()
        final_para.add_run("")

        # Row 3: Column headers (Step#, Action, Expected Result, Status)
        headers = ["Step#", "Action", "Expected Result", "Status"]
        for i, header in enumerate(headers):
            cell = table.rows[2].cells[i]
            cell.text = header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.bold = True
            run.font.size = Pt(12)  # +1 from base 11
            self._set_cell_background(cell, "a0c5e8")

        # Add data rows
        current_row = 3
        step_number = 1

        if self.test_case_steps:
            for step_data in self.test_case_steps:
                action_text = step_data.get('action', '')
                expected_text = step_data.get('expectedResult', '')
                status_text = step_data.get('status', '')

                # Check if this is a scenario (action contains "Scenario")
                if 'Scenario' in action_text or 'SCENARIO' in action_text.upper():
                    # Merge all 4 columns for scenario row
                    cell_scenario = table.rows[current_row].cells[0]
                    cell_scenario.merge(table.rows[current_row].cells[1])
                    cell_scenario.merge(table.rows[current_row].cells[2])
                    cell_scenario.merge(table.rows[current_row].cells[3])
                    cell_scenario.text = action_text
                    para = cell_scenario.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = para.runs[0]
                    run.font.bold = True
                else:
                    # Regular data row
                    table.rows[current_row].cells[0].text = str(step_number)
                    table.rows[current_row].cells[1].text = action_text
                    table.rows[current_row].cells[2].text = expected_text
                    table.rows[current_row].cells[3].text = status_text
                    step_number += 1

                current_row += 1

        # Last row: Screenshots (tiptap-editor-root-richeditor_395 content)
        screenshots_row_idx = total_rows - 1
        if screenshots_row_idx < len(table.rows):
            cell_screenshots = table.rows[screenshots_row_idx].cells[0]
            cell_screenshots.merge(table.rows[screenshots_row_idx].cells[1])
            cell_screenshots.merge(table.rows[screenshots_row_idx].cells[2])
            cell_screenshots.merge(table.rows[screenshots_row_idx].cells[3])

            # Add ordered content items (text and images in sequence)
            if self.screenshots_content_items:
                # Add a page break before screenshots section
                from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

                para = cell_screenshots.paragraphs[0]
                para.clear()
                para.add_run().add_break(WD_BREAK.PAGE)

                # Add "Screenshots/Additional Info:" as a header
                header_run = para.add_run("Screenshots/Additional Information:\n\n")
                header_run.font.bold = True

                # Set paragraph alignment to left
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                logging.info(f"[Word Export] ===== ADDING {len(self.screenshots_content_items)} CONTENT ITEMS IN ORDER =====")

                # Track image count for page breaks (2 images per page)
                image_count = 0

                # Process content items in order
                for idx, item in enumerate(self.screenshots_content_items):
                    try:
                        if item.get('type') == 'text':
                            # Add text content
                            text_content = item.get('content', '').strip()
                            if text_content:
                                text_run = para.add_run(text_content + "\n")
                                text_run.font.italic = True
                                logging.info(f"[Word Export] Added text ({len(text_content)} chars)")
                        elif item.get('type') == 'image':
                            # Add image
                            image_bytes = item.get('data')
                            if image_bytes:
                                import io
                                from docx.shared import Inches, Pt

                                # Add page break after every 2 images
                                if image_count > 0 and image_count % 2 == 0:
                                    para.add_run().add_break(WD_BREAK.PAGE)
                                    logging.info(f"[Word Export] Added page break after image {image_count}")

                                image_stream = io.BytesIO(image_bytes)

                                # Calculate size - enlarge images to fit ~2 per page
                                # Use 5.5 inches width for larger display
                                max_width = Inches(5.5)
                                img_width = item.get('width')
                                img_height = item.get('height')

                                if img_width and img_height:
                                    aspect_ratio = img_height / img_width
                                    final_width = max_width
                                    final_height = Inches(aspect_ratio * 5.5)
                                else:
                                    final_width = Inches(5.5)
                                    final_height = Inches(4)

                                # Add image
                                img_run = para.add_run()
                                img_run.add_picture(image_stream, width=final_width, height=final_height)

                                # Add newline after image
                                para.add_run("\n")
                                image_count += 1
                                logging.info(f"[Word Export] ✓✓✓ Added image {image_count} (enlarged, left-aligned)")
                    except Exception as e:
                        logging.error(f"[Word Export] ✗ Failed to add content item {idx + 1}: {e}", exc_info=True)
            else:
                cell_screenshots.text = "[No additional information available - Element may not have been loaded]"
                logging.warning("[Word Export] No screenshot content found, showing placeholder message")

    def _format_label_cell(self, cell, text):
        """Format a label cell with bold text and gray background"""
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        self._set_cell_background(cell, "d9d9d9")

    def _set_cell_background(self, cell, color_hex):
        """Set cell background color"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)

    def export_to_pdf(self):
        """Export data to PDF document"""
        logging.info("Export to PDF button clicked")

        # Log image availability
        if self.screenshots_images:
            logging.info(f"✓ {len(self.screenshots_images)} screenshot image(s) available for PDF export")
            for idx, img in enumerate(self.screenshots_images):
                logging.info(f"  Image {idx + 1}: type={img.get('type')}, data_size={len(img.get('data', b''))} bytes")
        else:
            logging.warning("⚠ No screenshot images available for PDF export")

        # Validate that Test Case Run Name is available
        test_case_run_name = self.test_case_run_name_entry.get()
        if not test_case_run_name:
            self.update_status("✗ Error: Test Script Name not loaded. Please load parameters first.", "error")
            return

        self.update_status("Exporting to PDF... Please wait.", "working")

        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            import os
            from datetime import datetime

            # Get current year
            current_year = datetime.now().year

            # Format script name
            script_name = self._format_script_name(test_case_run_name)

            # Collect form data
            data = self._collect_form_data()

            # Get save location with file dialog
            default_filename = f"{script_name}.pdf"
            filepath = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF Document", "*.pdf"), ("All Files", "*.*")],
                initialfile=default_filename,
                title="Save PDF Document"
            )

            if not filepath:
                self.update_status("Export cancelled by user.", "error")
                return

            # Create PDF
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            story = []

            # Add first page content
            first_page_elements = self._create_pdf_first_page(script_name)
            story.extend(first_page_elements)

            # Add page break for Page 2
            story.append(PageBreak())

            # Add Table 1: Author Details
            story.append(self._create_pdf_author_table(data))
            story.append(Spacer(1, 0.3*inch))

            # Add Table 2: Reviewer / Approver Roles
            story.append(self._create_pdf_reviewer_table(data))

            # Add page break for Page 3
            story.append(PageBreak())

            # Add Table 3: Execution Details
            story.append(self._create_pdf_execution_table(data))
            story.append(Spacer(1, 0.3*inch))

            # Add Table 4: Services Approval Sign-Off
            story.append(self._create_pdf_approval_table(data))

            # Add page break for Page 4
            story.append(PageBreak())

            # Verify images are still available before creating table
            if self.screenshots_images:
                logging.info(f"[PDF Export] Verified: {len(self.screenshots_images)} images ready for Table 5")
            else:
                logging.warning(f"[PDF Export] Warning: No images available for Table 5")

            # Add Table 5: Test Case Steps
            story.append(self._create_pdf_test_case_table(data))

            # Add ordered content items after the table (text and images in sequence)
            if self.screenshots_content_items:
                from reportlab.platypus import Image as RLImage, Paragraph, PageBreak
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_LEFT
                import io

                logging.info(f"[PDF Export] ===== ADDING {len(self.screenshots_content_items)} CONTENT ITEMS AFTER TABLE 5 =====")

                # Add page break before screenshots section
                story.append(PageBreak())

                # Add header with left alignment
                styles = getSampleStyleSheet()
                header_style = ParagraphStyle(
                    'ScreenshotHeader',
                    parent=styles['Heading2'],
                    alignment=TA_LEFT,
                    spaceAfter=12
                )
                story.append(Paragraph("<b>Screenshots/Additional Information:</b>", header_style))
                story.append(Spacer(1, 0.2*inch))

                # Style for text content - left aligned
                text_style = ParagraphStyle(
                    'ScreenshotText',
                    parent=styles['Normal'],
                    alignment=TA_LEFT,
                    leftIndent=0
                )

                # Track image count for page breaks (2 images per page)
                image_count = 0

                for idx, item in enumerate(self.screenshots_content_items):
                    try:
                        if item.get('type') == 'text':
                            # Add text content
                            text_content = item.get('content', '').strip()
                            if text_content:
                                # Clean HTML entities if present
                                text_content = text_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                text_para = Paragraph(f"<i>{text_content}</i>", text_style)
                                story.append(text_para)
                                story.append(Spacer(1, 0.15*inch))
                                logging.info(f"[PDF Export] Added text ({len(text_content)} chars)")
                        elif item.get('type') == 'image':
                            # Add image
                            image_bytes = item.get('data')
                            if image_bytes:
                                # Add page break after every 2 images
                                if image_count > 0 and image_count % 2 == 0:
                                    story.append(PageBreak())
                                    logging.info(f"[PDF Export] Added page break after image {image_count}")

                                image_stream = io.BytesIO(image_bytes)

                                # Calculate size - enlarge to fit ~2 images per page
                                # Using 5.5 inches width for larger display
                                max_width = 5.5 * inch
                                img_width = item.get('width')
                                img_height = item.get('height')

                                if img_width and img_height:
                                    aspect_ratio = img_height / img_width
                                    final_width = max_width
                                    final_height = aspect_ratio * max_width
                                else:
                                    final_width = 5.5 * inch
                                    final_height = 4 * inch

                                # Add image (left-aligned by default in reportlab)
                                rl_image = RLImage(image_stream, width=final_width, height=final_height)
                                story.append(rl_image)
                                story.append(Spacer(1, 0.3*inch))
                                image_count += 1
                                logging.info(f"[PDF Export] ✓✓✓ Added image {image_count} (enlarged, left-aligned)")
                    except Exception as e:
                        logging.error(f"[PDF Export] ✗ Failed to add content item {idx + 1}: {e}", exc_info=True)

            # Log export summary
            content_count = len(self.screenshots_content_items) if self.screenshots_content_items else 0
            if content_count > 0:
                logging.info(f"[PDF Export] ===== EXPORT COMPLETE: {content_count} content items included in document =====")
            else:
                logging.info(f"[PDF Export] ===== EXPORT COMPLETE: No content in document =====")

            # Build PDF with header/footer (no header/footer on first page)
            doc.build(story, onFirstPage=lambda c, d: None,
                     onLaterPages=lambda c, d: self._add_pdf_header_footer(c, d, script_name, current_year))

            logging.info(f"PDF document saved to: {filepath}")
            print(f"✓ PDF document exported to: {filepath}")

            # Update status with content count
            status_msg = f"✓ PDF document saved: {os.path.basename(filepath)}"
            if content_count > 0:
                status_msg += f" (with {content_count} content item(s))"
            self.update_status(status_msg, "success")

        except ImportError:
            error_msg = "reportlab library not installed. Please run: pip install reportlab"
            logging.error(error_msg)
            print(f"✗ {error_msg}")
            self.update_status(f"✗ {error_msg}", "error")
        except Exception as e:
            logging.error(f"Failed to export to PDF: {e}", exc_info=True)
            print(f"✗ Failed to export to PDF: {e}")
            self.update_status(f"✗ Failed to export to PDF: {str(e)}", "error")

    def _create_pdf_first_page(self, script_name):
        """Create first page elements for PDF"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from datetime import datetime
        import os

        styles = getSampleStyleSheet()
        elements = []

        # Get current date in format "22 Nov 2024"
        current_date = datetime.now().strftime("%d %b %Y")

        # Add spacing from top (approximately 2 lines)
        elements.append(Spacer(1, 0.3*inch))

        # Create table for title and "By..." line (2 rows, 1 column)
        # Create styles for table cells
        from reportlab.lib.enums import TA_CENTER

        # Script name: CENTER ALIGNED
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Normal'],
            fontSize=24,
            fontName='Helvetica',
            alignment=TA_CENTER,  # Center align script name
            textColor=colors.black,
            leftIndent=6,
            rightIndent=6,
            spaceBefore=6,
            leading=24 * 1.05  # Line spacing 1.05
        )

        # "By..." line: RIGHT ALIGNED
        from reportlab.lib.enums import TA_RIGHT
        by_style = ParagraphStyle(
            'ByStyle',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica',
            alignment=TA_RIGHT,  # Right align "By..." line
            textColor=colors.black,
            rightIndent=6,
            spaceBefore=6
        )

        # Create table data with Paragraph objects
        title_table_data = [
            [Paragraph(script_name, title_style)],
            [Paragraph("By Sycamore Informatics Inc.", by_style)]
        ]

        # Create table with specific row heights and column width
        title_table = Table(title_table_data, colWidths=[6.48*inch], rowHeights=[1.01*inch, 0.573*inch])
        title_table.setStyle(TableStyle([
            # Row 0: gray background
            ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#D9D9D9')),
            # Row 1: blue background
            ('BACKGROUND', (0, 1), (0, 1), colors.HexColor('#9FC5E8')),
            # Alignment
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            # Remove all padding
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))

        elements.append(title_table)

        # Add spacing before logo (one line space)
        elements.append(Spacer(1, 0.15*inch))

        # Add logo image (CENTER ALIGNED)
        logo_path = resource_path("logo.png")
        if os.path.exists(logo_path):
            # To center the logo, wrap it in a table
            logo = Image(logo_path, width=4.34*inch, height=2.20*inch)
            logo_table = Table([[logo]], colWidths=[6.5*inch])
            logo_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (0, 0), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]))
            elements.append(logo_table)

        # Add spacing before footer (6 centered empty paragraphs equivalent)
        elements.append(Spacer(1, 1.5*inch))

        # Add address on 3 separate lines (RIGHT ALIGNED)
        from reportlab.lib.enums import TA_RIGHT
        address_style = ParagraphStyle(
            'AddressStyle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Times-Roman',
            alignment=TA_RIGHT,  # Right align address and date
            textColor=colors.black
        )

        # Line 1: Address
        address_line1 = Paragraph("271 Waverley Oaks Rd #103", address_style)
        elements.append(address_line1)

        # Line 2: Company name
        address_line2 = Paragraph("Sycamore Informatics, Inc.", address_style)
        elements.append(address_line2)

        # Line 3: Date
        date_para = Paragraph(f"Date: {current_date}", address_style)
        elements.append(date_para)

        return elements

    def _add_pdf_header_footer(self, canvas, doc, script_name, current_year):
        """Add header and footer to PDF pages"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.units import inch
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT

        canvas.saveState()

        # Header - Create table with 2 columns (30% and 70%) to match docx layout
        page_width = letter[0] - 2*inch  # Total width minus margins

        # Create paragraph style for word-wrapped script name (right aligned, matches docx)
        script_name_style = ParagraphStyle(
            'ScriptNameStyle',
            fontName='Helvetica',
            fontSize=10,
            alignment=TA_RIGHT,
            wordWrap='CJK',  # Enable word wrap
        )

        # Create paragraph for script name to enable word wrap
        script_name_para = Paragraph(script_name, script_name_style)

        header_data = [["Sycamore Informatics", script_name_para]]
        header_table = Table(header_data, colWidths=[page_width * 0.30, page_width * 0.70])
        header_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica'),
            ('FONTSIZE', (0, 0), (0, 0), 10),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))

        # Draw header table
        header_table.wrapOn(canvas, page_width, inch)
        header_table.drawOn(canvas, inch, letter[1] - 0.6*inch)

        # Footer - Create table with 2 columns (85% and 15%)
        footer_text = f"© Sycamore Informatics {current_year}. Proprietary and Confidential Document, all rights reserved."
        page_text = f"Page {doc.page}"
        footer_data = [[footer_text, page_text]]
        footer_table = Table(footer_data, colWidths=[page_width * 0.85, page_width * 0.15])
        footer_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('WORDWRAP', (0, 0), (0, 0), False),
        ]))

        # Draw footer table
        footer_table.wrapOn(canvas, page_width, inch)
        footer_table.drawOn(canvas, inch, 0.5*inch)

        canvas.restoreState()

    def _create_pdf_author_table(self, data):
        """Create Author Details table for PDF"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        bold_style = styles['Normal'].clone('BoldStyle')
        bold_style.fontName = 'Helvetica-Bold'

        table_data = [
            ["Author Details"],
            [Paragraph("Author Name:", bold_style), Paragraph(data.get('Author Name', '[Not Set]'), normal_style)],
            [Paragraph("Author Date:", bold_style), Paragraph(data.get('Author Date', '[Not Set]'), normal_style)],
            [Paragraph("Last Updated By:", bold_style), Paragraph(data.get('Last Updated By', '[Not Set]'), normal_style)],
            [Paragraph("Last Updated Date:", bold_style), Paragraph(data.get('Last Updated Date', '[Not Set]'), normal_style)]
        ]

        table = Table(table_data, colWidths=[3*inch, 3*inch])
        table.setStyle(TableStyle([
            ('SPAN', (0, 0), (1, 0)),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#a0c5e8')),
            ('BACKGROUND', (0, 1), (0, 4), colors.HexColor('#d9d9d9')),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 13),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        return table

    def _create_pdf_reviewer_table(self, data):
        """Create Reviewer / Approver Roles table for PDF"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        bold_style = styles['Normal'].clone('BoldStyleReviewer')
        bold_style.fontName = 'Helvetica-Bold'
        bold_style.alignment = 1  # Center alignment

        # Create Paragraph objects for all columns to enable text wrapping
        responsibility_text = "To review and approve the script before the execution"

        table_data = [
            ["Reviewer / Approver"],
            [Paragraph("Name", bold_style), Paragraph("Role", bold_style), Paragraph("Responsibility", bold_style)],
            [Paragraph(data.get('Services Approver', '[Not Set]'), normal_style),
             Paragraph("Services Reviewer", normal_style),
             Paragraph(responsibility_text, normal_style)],
            [Paragraph(data.get('Product Management Approver', '[Not Set]'), normal_style),
             Paragraph("Product Management Reviewer", normal_style),
             Paragraph(responsibility_text, normal_style)],
            [Paragraph(data.get('Development Approver', '[Not Set]'), normal_style),
             Paragraph("Head of Engineering Reviewer", normal_style),
             Paragraph(responsibility_text, normal_style)]
        ]

        table = Table(table_data, colWidths=[1.8*inch, 1.8*inch, 2.4*inch])
        table.setStyle(TableStyle([
            ('SPAN', (0, 0), (2, 0)),
            ('ALIGN', (0, 0), (2, 0), 'CENTER'),
            ('BACKGROUND', (0, 0), (2, 0), colors.HexColor('#a0c5e8')),
            ('BACKGROUND', (0, 1), (2, 1), colors.HexColor('#d9d9d9')),
            ('FONTNAME', (0, 0), (2, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (2, 0), 13),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        return table

    def _create_pdf_execution_table(self, data):
        """Create Execution Details table for PDF"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        bold_style = styles['Normal'].clone('BoldStyleExecution')
        bold_style.fontName = 'Helvetica-Bold'

        table_data = [
            ["Execution Details"],
            [Paragraph("Release:", bold_style), Paragraph(data.get('Release', '[Not Set]'), normal_style)],
            [Paragraph("Executed By:", bold_style), Paragraph(data.get('Executor Name', '[Not Set]'), normal_style)],
            [Paragraph("Execution Start Date:", bold_style), Paragraph(data.get('Execution Start Date', '[Not Set]'), normal_style)],
            [Paragraph("Execution End Date:", bold_style), Paragraph(data.get('Execution End Date', '[Not Set]'), normal_style)],
            [Paragraph("Execution Status:", bold_style), Paragraph(data.get('Execution Status', '[Not Set]'), normal_style)],
            [Paragraph("Test Operating System:", bold_style), Paragraph(data.get('Test Operating System', '[Not Set]'), normal_style)],
            [Paragraph("Browser:", bold_style), Paragraph(data.get('Browser Name', '[Not Set]'), normal_style)],
            [Paragraph("URL:", bold_style), Paragraph(data.get('URL', '[Not Set]'), normal_style)],
            [Paragraph("Executor Signature:", bold_style), ""]
        ]

        table = Table(table_data, colWidths=[3*inch, 3*inch], rowHeights=[None]*9 + [0.7*inch])
        table.setStyle(TableStyle([
            ('SPAN', (0, 0), (1, 0)),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#a0c5e8')),
            ('BACKGROUND', (0, 1), (0, 8), colors.HexColor('#d9d9d9')),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 13),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        return table

    def _create_pdf_approval_table(self, data):
        """Create Services Approval Sign-Off table for PDF"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        bold_style = styles['Normal'].clone('BoldStyleApproval')
        bold_style.fontName = 'Helvetica-Bold'
        bold_centered = bold_style.clone('BoldCenteredApproval')
        bold_centered.alignment = 1  # Center alignment

        table_data = [
            ["Services Approval Sign-Off"],
            [Paragraph("Approver's Name", bold_centered), Paragraph("Signature", bold_centered)],
            [Paragraph(data.get('Services Approver', '[Not Set]'), normal_style), ""]
        ]

        table = Table(table_data, colWidths=[3*inch, 3*inch], rowHeights=[None, None, 0.7*inch])
        table.setStyle(TableStyle([
            ('SPAN', (0, 0), (1, 0)),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('BACKGROUND', (0, 0), (1, 0), colors.HexColor('#a0c5e8')),
            ('BACKGROUND', (0, 1), (1, 1), colors.HexColor('#d9d9d9')),
            ('BACKGROUND', (0, 2), (0, 2), colors.HexColor('#d9d9d9')),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 13),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        return table

    def _create_pdf_test_case_table(self, data):
        """Create Test Case Steps table for PDF (Table 5)"""
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import re

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        italic_style = styles['Normal'].clone('ItalicStyle')
        italic_style.fontName = 'Helvetica-Oblique'

        # Create bold styles for headers
        bold_center_style = ParagraphStyle(
            'BoldCenterStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=13,
            alignment=TA_CENTER
        )

        bold_center_12_style = ParagraphStyle(
            'BoldCenter12Style',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            alignment=TA_CENTER
        )

        bold_left_style = ParagraphStyle(
            'BoldLeftStyle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            alignment=TA_LEFT
        )

        # Build table data
        table_data = []

        # Row 1: Pre-Requisites header (merged)
        table_data.append([Paragraph("Pre-Requisites", bold_center_style), "", "", ""])

        # Row 2: Pre-Requisites content (merged, italic, with extra lines)
        prereq_text = self.prerequisites_text if self.prerequisites_text else "No prerequisites specified"
        # Add extra lines before and after, preserve line breaks in text
        prereq_formatted = prereq_text.replace('\n', '<br/>')
        prereq_para = Paragraph(f"<i><br/>{prereq_formatted}<br/></i>", normal_style)
        table_data.append([prereq_para, "", "", ""])

        # Row 3: Column headers
        table_data.append([
            Paragraph("Step#", bold_center_12_style),
            Paragraph("Action", bold_center_12_style),
            Paragraph("Expected Result", bold_center_12_style),
            Paragraph("Status", bold_center_12_style)
        ])

        # Add data rows
        step_number = 1
        if self.test_case_steps:
            for step_data in self.test_case_steps:
                action_text = step_data.get('action', '')
                expected_text = step_data.get('expectedResult', '')
                status_text = step_data.get('status', '')

                # Check if this is a scenario (merge all columns)
                if 'Scenario' in action_text or 'SCENARIO' in action_text.upper():
                    table_data.append([
                        Paragraph(action_text, bold_left_style),
                        "",
                        "",
                        ""
                    ])
                else:
                    # Regular data row
                    table_data.append([
                        Paragraph(str(step_number), normal_style),
                        Paragraph(action_text, normal_style),
                        Paragraph(expected_text, normal_style),
                        Paragraph(status_text, normal_style)
                    ])
                    step_number += 1

        # Last row: Screenshots (ordered content items)
        if self.screenshots_content_items:
            # Note in table that content is below
            num_images = len([item for item in self.screenshots_content_items if item.get('type') == 'image'])
            num_texts = len([item for item in self.screenshots_content_items if item.get('type') == 'text'])
            note_text = f"<b>Screenshots/Additional Information:</b><br/><br/>({num_texts} text block(s) and {num_images} image(s) below)"
            screenshots_para = Paragraph(note_text, normal_style)
            logging.info(f"[PDF Export] {len(self.screenshots_content_items)} content items will be added after table")
        else:
            screenshots_para = Paragraph("[No additional information available - Element may not have been loaded]", italic_style)
            logging.warning("[PDF Export] No screenshot content found, showing placeholder message")

        table_data.append([screenshots_para, "", "", ""])

        # Create table with column widths: 5%, 40%, 40%, 15%
        # Total width ~6.5": 0.325", 2.6", 2.6", 0.975"
        table = Table(table_data, colWidths=[0.325*inch, 2.6*inch, 2.6*inch, 0.975*inch])

        # Build style list for table
        style_list = [
            # Row 1: Pre-Requisites header - merge all columns, blue background
            ('SPAN', (0, 0), (3, 0)),
            ('BACKGROUND', (0, 0), (3, 0), colors.HexColor('#a0c5e8')),
            ('ALIGN', (0, 0), (3, 0), 'CENTER'),
            ('VALIGN', (0, 0), (3, 0), 'MIDDLE'),

            # Row 2: Pre-Requisites content - merge all columns
            ('SPAN', (0, 1), (3, 1)),

            # Row 3: Column headers - blue background
            ('BACKGROUND', (0, 2), (3, 2), colors.HexColor('#a0c5e8')),
            ('ALIGN', (0, 2), (3, 2), 'CENTER'),
            ('VALIGN', (0, 2), (3, 2), 'MIDDLE'),

            # Grid and padding for all cells
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]

        # Add span for scenario rows (rows with merged columns in data)
        # We need to identify which rows are scenarios
        row_idx = 3  # Start after headers
        if self.test_case_steps:
            for step_data in self.test_case_steps:
                action_text = step_data.get('action', '')
                if 'Scenario' in action_text or 'SCENARIO' in action_text.upper():
                    style_list.append(('SPAN', (0, row_idx), (3, row_idx)))
                row_idx += 1

        # Add span for last row (screenshots)
        style_list.append(('SPAN', (0, len(table_data) - 1), (3, len(table_data) - 1)))

        table.setStyle(TableStyle(style_list))

        return table

    def validate_required_fields(self):
        """Enable/disable Get Data from TP button based on required field validation"""
        run_id = self.id_entry.get().strip()
        username = self.user_name_entry.get().strip()
        password = self.password_entry.get().strip()

        # Enable button only if all three fields have values
        if run_id and username and password:
            self.load_page_button.configure(state="normal")
        else:
            self.load_page_button.configure(state="disabled")

    def update_status(self, message, status_type="info"):
        """Update status label with color coding
        status_type: 'info' (blue), 'success' (green), 'warning' (orange), 'error' (red), 'working' (yellow)
        """
        colors = {
            "info": "#3498db",
            "success": "#2ecc71",
            "warning": "#f39c12",
            "error": "#e74c3c",
            "working": "#f1c40f"
        }
        color = colors.get(status_type, "#3498db")
        self.status_label.configure(text=message, text_color=color)
        self.update_idletasks()  # Force UI update

    def on_closing(self):
        self.cmd_queue.put(("stop", None))
        self.playwright_thread.join(timeout=5)
        self.destroy()

def main():
    try:
        ctk.set_appearance_mode("Dark")
        ctk.set_default_color_theme("blue")
        app = TestCaseExporter()
        app.mainloop()
        print("Application closed.")
    except Exception as e:
        logging.error("Unhandled exception in main function", exc_info=True)
        raise

if __name__ == "__main__":
    main()
